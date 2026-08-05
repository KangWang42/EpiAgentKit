#!/usr/bin/env python3

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from generate_document_skill_showcase import (
    CN_HEADING,
    EN_FONT,
    add_body,
    add_caption,
    add_segments,
    add_three_line_table,
    add_title,
    configure_document,
    save_document,
    style_run,
)


ROOT = Path(__file__).resolve().parents[2]
OUTPUT_DIR = ROOT / "docs" / "showcase" / "graduate-opening-report"
FILL_MARKER = "【结构测试填充：以下 x 仅用于检验本节篇幅、分页与长文装配，不是正式研究内容】"
TEST_FILL_SCALE = 5


def pending(field):
    return f"【待补充：{field}】"


def citation(field):
    return f"【引文待核验：{field}】"


def test_fill(fill_chars, chunk_size=80):
    full_chunks, remainder = divmod(fill_chars, chunk_size)
    chunks = ["x" * chunk_size] * full_chunks
    if remainder:
        chunks.append("x" * remainder)
    return " ".join(chunks)


def test_text(opening, fill_chars, needs_citation=False):
    citation_marker = citation("本句所述外部事实或方法依据") if needs_citation else ""
    simulated_fill_chars = fill_chars * TEST_FILL_SCALE
    length_marker = f"【本节正文预算：约 {fill_chars} 个中文字符；测试填充：{simulated_fill_chars} 个 x】"
    return f"{opening}{citation_marker}{FILL_MARKER}{length_marker}{test_fill(simulated_fill_chars)}"


def configure_styles(document):
    configure_document(document)
    for style_name, size in (("Heading 1", 12), ("Heading 2", 11)):
        style = document.styles[style_name]
        style.font.name = EN_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
        fonts.set(qn("w:ascii"), EN_FONT)
        fonts.set(qn("w:hAnsi"), EN_FONT)
        fonts.set(qn("w:eastAsia"), CN_HEADING)
        style.paragraph_format.space_before = Pt(8 if style_name == "Heading 1" else 5)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.keep_with_next = True
    header = document.sections[0].header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_run(header.add_run("研究生学位论文开题报告 · 结构完整性测试模板"), size=8.5)


def add_heading(document, text, level=1):
    paragraph = document.add_paragraph(style=f"Heading {level}")
    paragraph.paragraph_format.keep_with_next = True
    add_segments(
        paragraph,
        [(text, {"size": 12 if level == 1 else 11, "bold": True, "east_asia": CN_HEADING})],
    )
    return paragraph


def add_toc(document):
    add_heading(document, "目录")
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(6)
    run = style_run(paragraph.add_run(), size=10)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = ' TOC \\o "1-2" \\h \\z \\u '
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "打开 Word 后更新目录字段"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, separate, placeholder, end))
    add_body(
        document,
        "目录字段仅用于测试 Word 装配；正式提交前应在目标模板中更新并核对页码。",
        first_line=False,
        size=9,
    )


def add_cover(document, report):
    for _ in range(3):
        document.add_paragraph()
    add_title(document, "硕士学位论文开题报告（结构完整性测试模板）")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(10)
    subtitle.paragraph_format.space_after = Pt(14)
    add_segments(subtitle, [(pending("论文题目"), {"size": 15, "bold": True, "east_asia": CN_HEADING})])
    note = document.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_segments(
        note,
        [("结构完整性测试模板：含显式占位符和测试填充，不得用于学院归档", {"size": 9})],
    )
    add_caption(document, "表 1　基本信息与测试边界", size=9)
    add_three_line_table(
        document,
        ["项目", "内容"],
        [
            ("学位层次", pending("硕士或博士")),
            ("专业与方向", pending("专业、研究方向")),
            ("作者与导师", pending("作者、学号、导师")),
            ("研究设计", report["design"]),
            ("模板状态", pending("学校或院系当前官方模板及版本")),
            ("篇幅校准", "按逐节预算模拟约 20–30 页 A4 完整报告"),
            ("测试用途", "验证完整章节、设计分支、篇幅、表格和 DOCX 装配"),
        ],
        [38, 118],
        aligns=["center", "left"],
        size=9.5,
    )
    document.add_page_break()
    add_toc(document)
    document.add_page_break()


def add_test_section(document, heading, opening, fill_chars, needs_citation=False, level=1):
    add_heading(document, heading, level=level)
    add_body(document, test_text(opening, fill_chars, needs_citation=needs_citation))


def add_table_block(document, caption_text, headers, rows, widths, aligns=None, size=8.8):
    add_caption(document, caption_text, size=8.8)
    table = add_three_line_table(document, headers, rows, widths, aligns=aligns, size=size)
    header_properties = table.rows[0]._tr.get_or_add_trPr()
    repeat_header = OxmlElement("w:tblHeader")
    repeat_header.set(qn("w:val"), "true")
    header_properties.append(repeat_header)
    return table


COMMON_SECTIONS = [
    ("一、立题背景及依据", "本节应说明研究问题、疾病或管理负担、已有证据和开展研究的必要性。", 4800, True),
    ("二、国内外研究现状与证据缺口", "本节应比较国内外证据的人群、设计、测量和结论差异并定位可回答的证据缺口。", 4800, True),
    ("三、研究意义", "本节应分别说明本研究可核查的理论意义、实践意义和适用边界。", 1200, True),
    ("四、研究目的、研究问题与研究假设", "本节应写明主要目的、次要目的、研究问题、研究假设和目标 estimand。", 1200, False),
    ("五、研究内容", "本节应按研究目的列出研究任务、先后依赖、输入、输出和停止条件。", 1500, False),
    ("六、研究方法", "本节应完整说明研究设计、研究对象、测量、样本量、质量控制、分析和伦理。", 500, False),
]


OBSERVATIONAL = {
    "filename": "observational-cohort.docx",
    "design": "前瞻性观察性队列研究",
    "method_sections": [
        ("6.1 研究设计、研究对象、研究场景和时间零点", "本节应定义目标人群、来源人群、研究场景、索引日期、基线和随访窗口。", 1000, False),
        ("6.2 纳入标准、排除标准、退出标准和失访定义", "本节应给出可操作判断的纳入、排除、退出、失访和分析集标准。", 900, False),
        ("6.3 暴露、比较组、协变量和混杂控制", "本节应定义暴露、参照组、时间顺序、候选协变量以及混杂、中介和碰撞变量的处理。", 1400, True),
        ("6.4 主要终点、次要终点及操作性定义", "本节应固定主要结局、次要结局、测量工具、单位、时间窗和判定规则。", 1400, True),
        ("6.5 样本量和精度依据", "本节应列出样本量或精度计算所需输入、效应尺度、把握度、失访率和来源。", 900, True),
        ("6.6 数据来源、测量工具、变量字典和数据管理", "本节应说明数据来源、授权、采集、编码、去标识化、访问权限、锁库和更正记录。", 1200, True),
        ("6.7 质量控制、选择偏倚、信息偏倚和缺失数据", "本节应规定培训、校准、复核、偏倚监测、缺失追踪和异常处理流程。", 1400, True),
        ("6.8 统计分析计划、主模型、效应量和置信区间", "本节应说明分析集、描述统计、主模型、调整变量、模型诊断和结果呈现。", 2000, True),
        ("6.9 敏感性分析、随访机制和非因果解释", "本节应预设缺失、失访、变量定义和模型假设的敏感性分析并限定为关联性解释。", 1100, True),
        ("6.10 伦理、知情同意、隐私保护、研究风险和应对措施", "本节应说明伦理状态、知情同意、隐私、数据安全、局限性、风险应对和修订条件。", 1100, True),
    ],
    "branch_table": (
        "表 4　暴露、参照组、协变量和时间顺序",
        ["项目", "操作性定义", "测量时间", "来源与核验"],
        [
            ("主要暴露", pending("暴露定义与编码"), pending("时间零点前测量"), citation("暴露依据")),
            ("参照组", pending("比较口径"), pending("与暴露相同窗口"), pending("锁定责任人")),
            ("混杂因素", pending("预设协变量及理由"), pending("暴露前测量"), citation("因果或方法依据")),
            ("中间事件", pending("随访中事件及处理"), pending("随访窗口"), pending("分析用途")),
        ],
        [28, 48, 32, 42],
        ["left", "left", "left", "left"],
    ),
    "route_nodes": ["对象筛查", "基线与暴露测量", "随访", "结局评估", "数据锁定", "关联分析与报告"],
    "analysis_label": "观察性关联分析、混杂调整和非因果解释",
}


RANDOMIZED = {
    "filename": "randomized-intervention.docx",
    "design": "平行组随机对照研究",
    "method_sections": [
        ("6.1 研究设计、研究对象、研究场景和时间点", "本节应定义目标人群、来源场景、招募流程、基线、随机时点和随访窗口。", 1000, False),
        ("6.2 纳入标准、排除标准、退出标准和失访定义", "本节应给出可操作判断的纳入、排除、退出、失访和分析集标准。", 900, False),
        ("6.3 随机序列、随机分配、分配隐藏和盲法", "本节应说明随机序列生成、分配比例、分配隐藏、盲法或开放标签理由及揭盲记录。", 1400, True),
        ("6.4 干预、对照组和实施忠实度", "本节应定义干预剂量、频次、持续时间、对照内容、共同措施、交叉和实施忠实度。", 1400, True),
        ("6.5 主要终点、次要终点和安全性结局", "本节应固定主要结局、次要结局、安全性结局、测量工具、时间窗和判定规则。", 1400, True),
        ("6.6 样本量和精度依据", "本节应列出样本量计算所需对照风险、目标差异、把握度、失访率和来源。", 900, True),
        ("6.7 数据来源、质量控制、方案偏离和缺失数据", "本节应规定实施记录、数据管理、培训校准、偏离分类、缺失追踪和异常处理。", 1400, True),
        ("6.8 意向性分析、符合方案分析和分析集", "本节应说明意向性分析、符合方案分析、主模型、效应量、置信区间和模型诊断。", 2000, True),
        ("6.9 敏感性分析、交叉和实施偏差", "本节应预设缺失、交叉、方案偏离、依从性和测量假设的敏感性分析。", 1100, True),
        ("6.10 伦理、安全性、不良事件和暂停条件", "本节应说明伦理、知情同意、隐私、安全性监测、不良事件报告和暂停修订条件。", 1100, True),
    ],
    "branch_table": (
        "表 4　随机化、干预、对照和实施忠实度",
        ["项目", "预设内容", "实施证据", "偏离与处理"],
        [
            ("随机序列", pending("生成方法和分组比例"), pending("保管和分配记录"), pending("错分处理")),
            ("干预组", pending("剂量、频次和持续时间"), pending("实施忠实度记录"), pending("未接受或交叉处理")),
            ("对照组", pending("对照内容和共同措施"), pending("实际接受记录"), pending("污染处理")),
            ("盲法", pending("对象、实施者和评估者状态"), pending("盲态或揭盲记录"), pending("开放标签偏差控制")),
        ],
        [28, 50, 38, 34],
        ["left", "left", "left", "left"],
    ),
    "route_nodes": ["对象筛查", "随机分组", "干预与对照实施", "随访和安全监测", "结局评估", "意向性分析与报告"],
    "analysis_label": "意向性分析、符合方案分析和安全性报告",
}


def common_method_tables(report):
    return [
        (
            "表 2　研究问题、estimand 与结局层级",
            ["层级", "研究问题", "estimand/结局", "对应分析"],
            [
                ("主要", pending("主要研究问题"), pending("人群、比较、结局和时间窗"), pending("主要分析")),
                ("次要", pending("次要研究问题"), pending("次要结局"), pending("次要分析")),
                ("探索性", pending("探索性问题"), pending("探索性结局"), pending("与主要分析分开")),
            ],
            [24, 42, 50, 34],
            ["center", "left", "left", "left"],
        ),
        (
            "表 3　纳入、排除、退出和失访判定",
            ["类别", "操作性标准", "判断时点", "记录位置"],
            [
                ("纳入", pending("逐条纳入标准"), pending("筛查时点"), pending("筛查表")),
                ("排除", pending("逐条排除标准"), pending("筛查时点"), pending("排除原因表")),
                ("退出", pending("撤回和终止定义"), pending("研究期间"), pending("退出记录")),
                ("失访", pending("联系次数和窗口"), pending("随访截止"), pending("失访表")),
            ],
            [25, 58, 32, 35],
            ["center", "left", "left", "left"],
        ),
        report["branch_table"],
        (
            "表 5　主要、次要和安全性终点定义",
            ["终点层级", "操作性定义与单位", "时间窗", "工具与缺失处理"],
            [
                ("主要终点", pending("主要终点定义"), pending("主要时间窗"), citation("测量或判定依据")),
                ("次要终点", pending("次要终点定义"), pending("次要时间窗"), citation("测量或判定依据")),
                ("安全性结局", pending("安全性结局定义"), pending("监测期间"), pending("报告与缺失规则")),
                ("探索性终点", pending("探索性终点"), pending("测量时间"), pending("与主要终点分开")),
            ],
            [28, 52, 28, 42],
            ["left", "left", "left", "left"],
        ),
        (
            "表 6　样本量或精度输入与核验责任",
            ["输入", "待填内容", "证据来源", "确认责任"],
            [
                ("主要结局参数", pending("比例、均值或事件率"), citation("先验研究或预实验"), pending("研究负责人")),
                ("效应尺度", pending("目标差异或精度"), citation("临床或方法依据"), pending("方法负责人")),
                ("统计参数", pending("显著性水平和把握度"), citation("方案依据"), pending("统计负责人")),
                ("设计修正", pending("失访、聚类或分层"), citation("实施依据"), pending("项目负责人")),
            ],
            [32, 42, 42, 34],
            ["left", "left", "left", "left"],
        ),
        (
            "表 7　预设统计分析计划",
            ["分析层级", "分析集与方法", "效应量和输出", "诊断或敏感性"],
            [
                ("主要", pending(report["analysis_label"]), pending("效应量、95% CI 和 P 值"), pending("模型诊断")),
                ("次要", pending("次要结局方法"), pending("估计与区间"), pending("多重性处理")),
                ("分层/交互", pending("预设因素"), pending("交互估计"), pending("探索性标记")),
                ("敏感性", pending("缺失和偏离假设"), pending("方向与区间比较"), pending("预设纳入条件")),
            ],
            [25, 48, 42, 35],
            ["left", "left", "left", "left"],
        ),
        (
            "表 8　质量控制、偏倚和风险应对",
            ["风险", "监测证据", "控制或应对", "停止/修订条件"],
            [
                ("对象进入", pending("筛查、选择或分配记录"), pending("连续记录和复核"), pending("流程偏离阈值")),
                ("测量质量", pending("培训、校准和重复测量"), pending("回到原始来源核对"), pending("测量失效条件")),
                ("随访/实施", pending("失访或忠实度记录"), pending("追踪和偏离处理"), pending("可解释性受损条件")),
                ("数据与合规", pending("权限、锁库和伦理状态"), pending("去标识化和审查"), pending("暂停或报告条件")),
            ],
            [28, 42, 45, 35],
            ["left", "left", "left", "left"],
        ),
    ]


def add_later_sections(document, report):
    add_test_section(document, "七、技术路线与研究流程", "本节应把对象进入、资料或干预、质量控制、结局评估、数据锁定、分析和产出连接起来。", 800)
    route_rows = [
        (node, pending(f"{node}的输入与主要动作"), pending(f"{node}的关键输出"), pending(f"{node}的停止或修订条件"))
        for node in report["route_nodes"]
    ]
    add_table_block(
        document,
        "表 9　技术路线与停止条件",
        ["节点", "主要任务", "关键输出", "停止/修订条件"],
        route_rows,
        [28, 52, 40, 30],
        ["left", "left", "left", "left"],
    )
    add_test_section(document, "八、预期结果及预期成果", "本节应列出预期形成的结果类型、表格、图件、数据库、论文部件和解释边界。", 700)
    add_test_section(document, "九、研究亮点及创新性", "本节应依据真实文献和研究方案比较设计、数据、测量、方法、实施或应用层面的可核查差异。", 900, needs_citation=True)
    add_test_section(document, "十、可行性分析与风险应对", "本节应区分已经具备、待确认、需要申请和存在风险的条件并给出应对责任。", 1000)
    add_table_block(
        document,
        "表 10　可行性条件与风险应对",
        ["维度", "当前状态", "证据或待确认事项", "应对与责任"],
        [
            ("数据/样本", pending("已具备、待确认或需申请"), pending("来源、授权和进入路径"), pending("责任人和截止时间")),
            ("技术与人员", pending("能力和培训状态"), pending("设备、人员和统计能力"), pending("培训或补充方案")),
            ("伦理与隐私", pending("伦理和授权状态"), pending("审批、同意和访问权限"), pending("伦理责任人")),
            ("时间与经费", pending("周期和预算状态"), pending("关键窗口和资源"), pending("调整条件")),
            ("主要风险", pending("风险等级"), pending("监测指标"), pending("应对和停止条件")),
        ],
        [28, 35, 52, 35],
        ["left", "left", "left", "left"],
    )
    add_test_section(document, "十一、研究工作计划", "本节应按真实研究周期排列证据核验、伦理、工具准备、对象进入、实施随访、锁库、分析、写作和提交。", 500)
    add_table_block(
        document,
        "表 11　研究工作计划",
        ["阶段", "计划时间", "主要工作", "阶段性产出"],
        [
            ("第一阶段", pending("起止月份"), pending("文献、方案和伦理"), pending("阶段产出")),
            ("第二阶段", pending("起止月份"), pending("工具、培训和对象进入"), pending("阶段产出")),
            ("第三阶段", pending("起止月份"), pending("资料收集、干预或随访"), pending("阶段产出")),
            ("第四阶段", pending("起止月份"), pending("锁库、分析和结果复核"), pending("阶段产出")),
            ("第五阶段", pending("起止月份"), pending("论文、预答辩和正式提交"), pending("阶段产出")),
        ],
        [25, 32, 60, 33],
        ["center", "center", "left", "left"],
    )


def add_references(document):
    add_test_section(document, "十二、参考文献", "本节应列出与正文引文一一对应且完成身份、内容和适用性核验的真实来源。", 500, needs_citation=True)
    reference_needs = [
        "研究问题与目标人群的背景来源",
        "疾病负担或应用问题的权威来源",
        "国内研究现状的代表性来源",
        "国外研究现状的代表性来源",
        "研究设计与偏倚控制的方法依据",
        "测量工具、终点或量表依据",
        "样本量或统计分析的方法依据",
        "伦理、隐私或报告规范依据",
    ]
    for index, need in enumerate(reference_needs, start=1):
        add_body(document, f"[{index}] {citation(need)}", first_line=False, size=9.2, after=2)


def add_appendix(document):
    add_test_section(document, "十三、附件、导师意见与专家评议表单", "本节应按官方模板保留必要附件、导师意见、专家评议、开题考核、签字和日期空位。", 500)
    add_table_block(
        document,
        "表 12　附件与签字栏清单",
        ["附件/表单", "用途", "测试状态"],
        [
            ("附件 1 变量与终点字典", "统一变量、单位、时间点和操作性定义", pending("正式内容")),
            ("附件 2 技术路线和研究流程", "核对对象、输入、处理、输出和停止条件", pending("正式图表")),
            ("附件 3 研究工具或实施材料", "放置量表、问卷、材料或访谈提纲", pending("实际附件")),
            ("导师意见", "保留官方意见正文、签字和日期空位", "不得代填"),
            ("专家评议与考核", "保留评议、结论、签字和日期字段", "不得代填"),
        ],
        [42, 74, 34],
        aligns=["left", "left", "center"],
    )
    add_body(
        document,
        "本文件只验证结构、篇幅和 DOCX 装配，全部 x、待补充项和引文待核验项必须在正式归档前由真实研究材料替换并通过 archive 模式检查。",
        first_line=False,
        size=9.5,
    )


def build_report(report):
    document = Document()
    configure_styles(document)
    document.core_properties.title = "研究生学位论文开题报告结构完整性测试模板"
    document.core_properties.subject = report["design"]
    document.core_properties.author = "结构测试"
    add_cover(document, report)
    for heading, opening, fill_chars, needs_citation in COMMON_SECTIONS:
        add_test_section(document, heading, opening, fill_chars, needs_citation=needs_citation)
        if heading == "六、研究方法":
            for subheading, subopening, subfill, subcitation in report["method_sections"]:
                add_test_section(document, subheading, subopening, subfill, needs_citation=subcitation, level=2)
            for table in common_method_tables(report):
                add_table_block(document, *table)
    add_later_sections(document, report)
    add_references(document)
    add_appendix(document)
    output_path = OUTPUT_DIR / report["filename"]
    save_document(document, output_path)
    return output_path


def main():
    for report in (OBSERVATIONAL, RANDOMIZED):
        path = build_report(report)
        print(f"已生成 {path.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
