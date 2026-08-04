#!/usr/bin/env python3

import csv
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


DEMO_DIR = Path(__file__).resolve().parent
FIGURE_DIR = DEMO_DIR / "output" / "publication-figures"
OUTPUT_DIR = DEMO_DIR / "output" / "academic-publishing"
RESULTS_PATH = FIGURE_DIR / "survival-demo-results.csv"
COHORT_PATH = FIGURE_DIR / "manuscript-cohort-summary.csv"
DIAGNOSTICS_PATH = FIGURE_DIR / "model-diagnostics.csv"
FOREST_RESULTS_PATH = FIGURE_DIR / "cox-forest-results.csv"
SURVIVAL_FIGURE_PATH = FIGURE_DIR / "adjusted-survival-paper.png"
FOREST_FIGURE_PATH = FIGURE_DIR / "cox-forest.png"
ZH_DOCX_PATH = OUTPUT_DIR / "manuscript-preview-zh.docx"
EN_DOCX_PATH = OUTPUT_DIR / "manuscript-preview-en.docx"


def style_run(run, size=10.5, bold=False, italic=False, east_asia="宋体"):
    run.bold = bold
    run.italic = italic
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(23, 32, 51)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    return run


def add_run(paragraph, text, **formatting):
    return style_run(paragraph.add_run(text), **formatting)


def format_body(paragraph, size=10.5, after=4, first_line=True):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(after)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Pt(size * 2)
    return paragraph


def add_body(document, text, size=10.5, east_asia="宋体", after=4, first_line=True):
    if east_asia != "Times New Roman" and size >= 10 and first_line and len(text) < 150:
        raise ValueError(f"中文正文段落过短（{len(text)} 字符）：{text[:24]}")
    paragraph = format_body(
        document.add_paragraph(), size=size, after=after, first_line=first_line
    )
    add_run(paragraph, text, size=size, east_asia=east_asia)
    return paragraph


def add_heading(document, text, east_asia="黑体", size=12, level=1):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    add_run(paragraph, text, size=size, bold=True, east_asia=east_asia)
    return paragraph


def configure_page(document, margin_mm=21):
    for section in document.sections:
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.top_margin = Mm(margin_mm)
        section.bottom_margin = Mm(margin_mm)
        section.left_margin = Mm(22)
        section.right_margin = Mm(22)
        section.header_distance = Mm(9)
        section.footer_distance = Mm(9)


def configure_normal_style(document, size=10.5, east_asia="宋体"):
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(size)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def set_zoom(document, percent=100):
    zoom = document.settings._element.xpath("./w:zoom")
    if zoom:
        zoom[0].set(qn("w:percent"), str(percent))


def add_page_number(section, east_asia="宋体"):
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = style_run(paragraph.add_run(), size=9, east_asia=east_asia)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


def set_cell_border(cell, edge, value="nil", size=0, color="auto"):
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        later_properties = {
            qn(f"w:{name}")
            for name in (
                "shd",
                "noWrap",
                "tcMar",
                "textDirection",
                "tcFitText",
                "vAlign",
                "hideMark",
            )
        }
        insertion_index = len(properties)
        for index, child in enumerate(properties):
            if child.tag in later_properties:
                insertion_index = index
                break
        properties.insert(insertion_index, borders)
    border = borders.find(qn(f"w:{edge}"))
    if border is None:
        border = OxmlElement(f"w:{edge}")
        borders.append(border)
    border.set(qn("w:val"), value)
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), color)


def set_three_line_table_borders(table):
    for row in table.rows:
        for cell in row.cells:
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                set_cell_border(cell, edge)
    for cell in table.rows[0].cells:
        set_cell_border(cell, "top", value="single", size=8, color="000000")
        set_cell_border(cell, "bottom", value="single", size=6, color="000000")
    for cell in table.rows[-1].cells:
        set_cell_border(cell, "bottom", value="single", size=8, color="000000")


def set_cell_text(cell, text, size=8.5, bold=False, east_asia="宋体"):
    cell.text = ""
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    add_run(
        paragraph,
        str(text),
        size=size,
        bold=bold,
        east_asia=east_asia,
    )


def add_table(document, headers, rows, widths_mm, east_asia="宋体", size=8.5):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for index, (header, width) in enumerate(zip(headers, widths_mm)):
        table.columns[index].width = Mm(width)
        set_cell_text(table.rows[0].cells[index], header, size=size, bold=True, east_asia=east_asia)
    for row_values in rows:
        cells = table.add_row().cells
        for column_index, (value, width) in enumerate(zip(row_values, widths_mm)):
            cells[column_index].width = Mm(width)
            set_cell_text(cells[column_index], value, size=size, east_asia=east_asia)
    set_three_line_table_borders(table)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_caption(document, text, east_asia="宋体", size=9):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(2)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_with_next = False
    add_run(paragraph, text, size=size, east_asia=east_asia)


def add_figure(document, path, caption, width_mm=154, east_asia="宋体"):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.add_run().add_picture(str(path), width=Mm(width_mm))
    add_caption(document, caption, east_asia=east_asia)


def add_reference(document, text, east_asia="宋体", size=9):
    paragraph = format_body(
        document.add_paragraph(), size=size, after=2, first_line=False
    )
    paragraph.paragraph_format.left_indent = Pt(14)
    paragraph.paragraph_format.first_line_indent = Pt(-14)
    add_run(paragraph, text, size=size, east_asia=east_asia)


def read_csv_rows(path):
    with path.open(encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def format_p_value(value):
    numeric = float(value)
    return "<0.001" if numeric < 0.001 else f"={numeric:.3f}"


def build_chinese_manuscript(result, cohort_rows, diagnostic_rows, forest_rows):
    n = int(result["n"])
    events = int(result["events"])
    hazard_ratio = float(result["hazard_ratio"])
    ci_lower = float(result["ci_lower"])
    ci_upper = float(result["ci_upper"])
    treatment_p = format_p_value(result["p_value"])
    diagnostics = {row["statistic"]: float(row["value"]) for row in diagnostic_rows}
    model_terms = {row["term"]: row for row in forest_rows}

    document = Document()
    configure_page(document)
    configure_normal_style(document)
    add_page_number(document.sections[0])

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    add_run(
        title,
        "固定模拟队列中治疗方案与 36 个月无事件生存的关联",
        size=16,
        bold=True,
        east_asia="黑体",
    )

    abstract = format_body(document.add_paragraph(), after=4, first_line=False)
    add_run(abstract, "摘要　", bold=True, east_asia="黑体")
    add_run(
        abstract,
        "目的：在分析流程、结果取数和论文表达均可复核的条件下，评估固定模拟队列中强化治疗方案与 36 个月无事件生存的关联。"
        "方法：采用预设随机种子生成队列资料，行政删失时间为 36 个月。以 Cox 比例风险模型估计治疗方案与结局发生风险的关联，并调整年龄、性别、疾病分期和标准化生物标志物。"
        f"结果：共纳入 {n:,} 名模拟研究对象，观察到 {events:,} 个结局事件。调整协变量后，强化方案组的结局发生风险低于常规方案组（风险比 = {hazard_ratio:.2f}，95% 置信区间：{ci_lower:.2f}～{ci_upper:.2f}）。"
        "比例风险全局检验未见明显偏离。结论：在该模拟数据结构和预设模型下，强化方案与较低的结局发生风险相关。该结果仅用于验证可复核工作流，不构成真实临床证据。",
    )
    keywords = format_body(document.add_paragraph(), after=6, first_line=False)
    add_run(keywords, "关键词　", bold=True, east_asia="黑体")
    add_run(keywords, "模拟队列；无事件生存；Cox 比例风险模型；结果追溯")

    add_heading(document, "1　引言")
    add_body(
        document,
        "时间到事件结局同时包含结局是否发生及其发生时间。在存在删失的条件下，直接比较粗事件比例会丢失随访时间信息，也可能使组间可比性受随访长度影响。Kaplan–Meier 方法可用于描述删失条件下的生存分布，Cox 比例风险模型则允许在协变量条件下估计相对风险[1–2]。观察性资料中的治疗选择往往与年龄、疾病严重程度和基线生物学特征有关，所以未调整比较可能同时反映治疗方案、基线构成和随访过程。研究报告因而需同时说明时间零点、分析集、删失规则、协变量调整、参照水平和不确定性，避免将模型关联改写为已经证实的治疗因果效应。",
    )
    add_body(
        document,
        "除模型选择外，数字的来源也决定研究结果能否被复核。数据生成、模型估计、机器可读结果、图表和正文如果分别手工维护，修订时容易出现方向、精度或版本不一致。可复核的作者稿应使研究目的、分析对象、表图和论断使用同一组已验证结果，并在数字变更时回到产生估计的分析。本研究使用固定模拟队列，完整展示从数据产生、协变量处理、模型运行、假设检查到文稿表达的一致链路。研究问题限定为强化方案与 36 个月无事件生存的关联，不评价治疗机制，也不将模拟估计作为真实临床证据。",
    )

    add_heading(document, "2　资料与方法")
    add_heading(document, "2.1　研究设计、暴露与结局", size=11, level=2)
    add_body(
        document,
        "本研究为基于固定随机种子的模拟队列分析。研究对象在统一时间零点进入队列，最长随访 36 个月。暴露为治疗方案，分为常规方案与强化方案；治疗分配概率由年龄、疾病分期和生物标志物共同决定，以形成符合观察性研究语境的组间差异。主要结局为 36 个月内首次发生的模拟事件，事件时间由预设基线风险和协变量线性预测值生成，未发生事件者在 36 个月时行政删失。预设协变量包括年龄、性别、疾病分期和基线生物标志物。数据不对应真实患者、机构或干预，报告结构参照 STROBE 对观察性研究透明报告的核心要求，但不将报告指南作为研究质量评价工具[4–5]。",
    )
    add_heading(document, "2.2　统计分析与质量控制", size=11, level=2)
    add_body(
        document,
        "按治疗方案描述模拟队列的基线特征和事件数。主要分析采用 Cox 比例风险模型，治疗方案以常规方案为参照，年龄按每增加 10 岁计入，性别以女性为参照，疾病分期以 I 期为参照，生物标志物按一个标准差计量。报告风险比（hazard ratio，HR）及其 95% 置信区间。以调整后生存曲线表示在相同协变量取值下两组的无事件生存概率，同时给出 95% 置信区间和 0、12、24、36 个月的风险集。以 Schoenfeld 残差检验比例风险假设，并报告全局检验结果[3]；一致性指数仅作为排序区分能力的补充描述，不用于判定治疗的临床价值。",
    )
    add_body(
        document,
        "模拟数据由固定随机种子生成，分析脚本从同一数据文件重建分类水平、构造预设连续变量、拟合模型并导出机器可读结果。表格、图件和正文使用同一次模型导出值，数字修正必须回到产生结果的分析，不在成品中单独改写。质量控制按数据、模型和成品三个层次完成：核对总样本量与事件数，检查模型项、计量单位和参照水平，再比较正文、表格和图件中的估计方向、置信区间与显示精度。比例风险检验和一致性指数同步导出，避免把模型诊断与主要效应估计分离成不同版本。",
    )

    add_heading(document, "3　结果")
    add_heading(document, "3.1　研究对象与基线特征", size=11, level=2)
    regular, intensive = cohort_rows
    add_body(
        document,
        f"共纳入 {n:,} 名模拟研究对象，其中常规方案组 {int(regular['n']):,} 人，强化方案组 {int(intensive['n']):,} 人。"
        f"随访期内分别观察到 {int(regular['events']):,} 个和 {int(intensive['events']):,} 个结局事件。"
        f"常规方案组与强化方案组的平均年龄分别为 {float(regular['age_mean']):.1f} 岁和 {float(intensive['age_mean']):.1f} 岁，男性比例分别为 {float(regular['male_pct']):.1f}% 和 {float(intensive['male_pct']):.1f}%。"
        f"III 期比例在两组中分别为 {float(regular['stage_iii_pct']):.1f}% 和 {float(intensive['stage_iii_pct']):.1f}%，生物标志物均值分别为 {float(regular['biomarker_mean']):.2f} 和 {float(intensive['biomarker_mean']):.2f}。"
        "年龄和性别构成整体接近，但疾病分期和生物标志物分布存在一定差异，因此未根据粗事件比例作效果判断，而是在主要模型中按预设方案同时调整这些协变量（表 1）。",
    )
    add_caption(document, "表 1　按治疗方案分组的模拟队列特征")
    table_rows = []
    for row in cohort_rows:
        table_rows.append(
            (
                row["group"],
                f"{int(row['n'])}",
                f"{int(row['events'])}",
                f"{float(row['age_mean']):.1f} ± {float(row['age_sd']):.1f}",
                f"{int(row['male_n'])} ({float(row['male_pct']):.1f})",
                f"{int(row['stage_iii_n'])} ({float(row['stage_iii_pct']):.1f})",
                f"{float(row['biomarker_mean']):.2f} ± {float(row['biomarker_sd']):.2f}",
            )
        )
    add_table(
        document,
        ("治疗方案", "人数", "事件数", "年龄（岁）", "男性，n (%)", "III 期，n (%)", "生物标志物"),
        table_rows,
        (27, 16, 18, 28, 27, 25, 29),
    )
    add_body(
        document,
        "注：连续变量以均数 ± 标准差表示；分类变量以人数（构成比）表示。全部数据为固定随机种子生成的模拟资料。",
        size=9,
        first_line=False,
    )

    add_heading(document, "3.2　治疗方案与无事件生存", size=11, level=2)
    add_body(
        document,
        f"调整年龄、性别、疾病分期和标准化生物标志物后，强化方案组的结局发生风险低于常规方案组（HR = {hazard_ratio:.2f}，95% 置信区间：{ci_lower:.2f}～{ci_upper:.2f}，P {treatment_p}）。"
        "在设定相同年龄、性别、疾病分期和生物标志物水平时，强化方案组的调整后无事件生存曲线在随访期内整体高于常规方案组，两条曲线的估计方向与 Cox 模型的 HR 一致（图 1）。图中同时显示置信区间和各时点风险集，使曲线差异能在样本量逐渐减少的条件下解读。该差异表示预设模型条件下的统计关联，不支持对真实治疗效果作出推断。",
    )
    add_figure(
        document,
        SURVIVAL_FIGURE_PATH,
        "图 1　两种治疗方案的调整后无事件生存曲线及 95% 置信区间",
        width_mm=150,
    )

    add_heading(document, "3.3　多变量模型与假设检查", size=11, level=2)
    age_term = model_terms["age_10"]
    male_term = model_terms["sex男性"]
    stage_ii = model_terms["stageII"]
    stage_iii = model_terms["stageIII"]
    biomarker = model_terms["biomarker_sd"]
    add_body(
        document,
        f"在同一多变量模型中，年龄每增加 10 岁的 HR 为 {float(age_term['estimate']):.2f}（95% 置信区间：{float(age_term['conf_low']):.2f}～{float(age_term['conf_high']):.2f}）；"
        f"男性与女性比较的 HR 为 {float(male_term['estimate']):.2f}（95% 置信区间：{float(male_term['conf_low']):.2f}～{float(male_term['conf_high']):.2f}）。"
        f"与 I 期比较，II 期和 III 期的 HR 分别为 {float(stage_ii['estimate']):.2f} 和 {float(stage_iii['estimate']):.2f}；"
        f"生物标志物每增加一个标准差的 HR 为 {float(biomarker['estimate']):.2f}。各项估计的 95% 置信区间均未跨越 1，但这些协变量的结果用于说明调整模型的构成和方向，不作为额外的临床发现。全部模型项的效应估计、参照水平及不确定性见图 2。",
    )
    add_body(
        document,
        f"Schoenfeld 残差全局检验 P = {diagnostics['global_ph_test_p']:.3f}，在本模拟样本与检验效能下未见明显的比例风险假设偏离。"
        f"模型一致性指数为 {diagnostics['concordance']:.3f}（标准误 = {diagnostics['concordance_se']:.3f}），表明在该模拟数据中具有中等程度的排序区分能力。比例风险检验反映模型中协变量效应是否随时间系统变化，一致性指数则反映个体风险排序，两者不能相互替代。这些诊断结果与主要 HR 一并报告，便于读者判断效应表达的可解释性；但一致性指数不反映预测概率的校准，也不等同于临床决策效用。",
    )
    add_figure(
        document,
        FOREST_FIGURE_PATH,
        "图 2　多变量 Cox 比例风险模型的效应估计",
        width_mm=156,
    )

    add_heading(document, "4　讨论")
    add_body(
        document,
        f"本模拟队列分析显示，在调整预设基线特征后，强化方案与较低的 36 个月结局发生风险相关，效应估计为 HR {hazard_ratio:.2f}。调整后生存曲线、森林图和正文使用同一次模型结果，估计方向、置信区间、参照水平和显示精度相互一致。这种一致性使读者能从论文中的主要论断回到具体模型项，再核对曲线和风险集是否支持同一方向的解释。因此，本例验证的是从分析对象到发表级图件和论文文本的可追溯装配，而不是对某项真实干预效果的评价。模拟数据的估计大小只反映预设参数和随机变异，不应与任何疾病、药物或人群的真实效果建立对应。",
    )
    add_body(
        document,
        "主要效应的解释必须结合模型尺度。HR 是给定协变量条件下的相对即时风险，不直接等于 36 个月的绝对风险差，也不能单独表示个体层面的预后。图 1 通过设定共同协变量取值展示两组生存概率，与 HR 提供的相对效应信息互补；风险集则提示后期曲线估计所依据的可观察对象数逐渐减少，因而不应只根据曲线间的视觉距离判断效应强度。实际研究还应根据预先定义的 estimand 报告指定时点的绝对风险、生存率差异或限制平均生存时间，并说明这些绝对效应与相对即时风险的不同含义。只有当时间零点、随访范围和分析集一致时，这些指标才能共同回答同一个研究问题。",
    )
    add_body(
        document,
        "本例中治疗分配概率受基线特征影响，因此设计上存在混杂。多变量回归只能调整已测量且函数形式合理的协变量，无法排除未测量混杂、测量误差或选择偏倚。即使在真实观察性队列中观察到类似结果，也应使用“相关”或“关联”表述，不应将其改写为已证明的因果效应。比例风险全局检验未见明显偏离，但不显著检验不能证明假设必然成立；在正式研究中，还需结合分时段效应、残差图和临床过程判断假设的合理性。同样，一致性指数只表示排序区分，无法替代校准、决策曲线或对预测模型临床使用方式的独立评价[6]。这些边界决定了模型诊断应与效应估计一同解释，不能把某一项检验通过当作全部科学假设已被验证。",
    )

    add_heading(document, "5　局限性")
    add_body(
        document,
        "本研究存在明确的适用边界。首先，数据由完全知情的随机过程生成，不包含真实队列中常见的缺失、测量误差、竞争风险、访视不规则和治疗随时变化，因此无法演示这些问题对分析集、估计偏倚和不确定性的影响。其次，模型与数据生成机制相对匹配，不能代表错误函数形式、非线性关系或重要交互作用被遗漏时的表现。再次，本文没有估计平均治疗效应、处理时变混杂，也没有进行外部验证或临床决策分析。最后，固定随机种子保证示例结果可重现，却不能代替真实数据来源、测量质量和偏倚机制的评估。这些局限共同决定了结果只能用于检查工作流和表达一致性。",
    )

    add_heading(document, "6　结论")
    add_body(
        document,
        "在固定模拟队列和预设 Cox 比例风险模型下，强化方案与较低的 36 个月结局发生风险相关。表格、调整后生存曲线、森林图和正文数字均来自同一次已验证分析，主要效应、参照水平、置信区间和模型诊断能够相互核对。该示例说明了完整作者稿应如何在保持证据边界的同时实现结果追溯，但其中的治疗方案、效应估计和生存差异均只属于模拟情景。任何真实研究都需根据实际人群、暴露定义、数据质量和偏倚结构重新完成分析，不得将本文的模拟估计外推为医学结论。",
    )

    add_heading(document, "声明")
    add_body(
        document,
        "数据与伦理：本文仅使用计算机生成的模拟数据，不包含真实个人资料，因此不涉及研究对象同意或伦理审查。利益冲突与资助：本示例未设定作者、机构或资助来源，不以虚构信息填充投稿声明。数据可得性：模拟数据、分析脚本和模型导出结果随示例仓库公开。",
        first_line=False,
    )

    add_heading(document, "参考文献")
    references = (
        "[1] Kaplan EL, Meier P. Nonparametric estimation from incomplete observations. Journal of the American Statistical Association. 1958;53:457–481. doi:10.1080/01621459.1958.10501452.",
        "[2] Cox DR. Regression models and life-tables. Journal of the Royal Statistical Society: Series B. 1972;34:187–220. doi:10.1111/j.2517-6161.1972.tb00899.x.",
        "[3] Grambsch PM, Therneau TM. Proportional hazards tests and diagnostics based on weighted residuals. Biometrika. 1994;81:515–526. doi:10.1093/biomet/81.3.515.",
        "[4] von Elm E, Altman DG, Egger M, et al. The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement. PLoS Medicine. 2007;4:e296. doi:10.1371/journal.pmed.0040296.",
        "[5] Vandenbroucke JP, von Elm E, Altman DG, et al. Strengthening the Reporting of Observational Studies in Epidemiology (STROBE): explanation and elaboration. PLoS Medicine. 2007;4:e297. doi:10.1371/journal.pmed.0040297.",
        "[6] Steyerberg EW, Vickers AJ, Cook NR, et al. Assessing the performance of prediction models: a framework for traditional and novel measures. Epidemiology. 2010;21:128–138. doi:10.1097/EDE.0b013e3181c30fb2.",
    )
    for reference in references:
        add_reference(document, reference)

    set_zoom(document)
    document.core_properties.title = "固定模拟队列中治疗方案与 36 个月无事件生存的关联"
    document.core_properties.subject = "中文观察性分析完整作者稿示例"
    document.save(ZH_DOCX_PATH)


def build_english_manuscript():
    document = Document()
    configure_page(document)
    configure_normal_style(document, size=10.5, east_asia="Times New Roman")
    add_page_number(document.sections[0], east_asia="Times New Roman")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(10)
    add_run(
        title,
        "A Reproducible Framework for Interpreting Calibration Slopes in External Model Evaluation",
        size=16,
        bold=True,
        east_asia="Times New Roman",
    )

    abstract = format_body(document.add_paragraph(), after=4, first_line=False)
    add_run(
        abstract,
        "Abstract  ",
        bold=True,
        east_asia="Times New Roman",
    )
    add_run(
        abstract,
        "Calibration slopes are frequently reduced to a binary adequacy label even though their meaning depends on the definition of the linear predictor, the evaluation sample, and the fitted recalibration model. This methodological article sets out a reproducible assessment framework that separates overall calibration, spread of predictions, local agreement, discrimination, and decision consequences. The framework requires investigators to freeze the original prediction rule, declare the target population and prediction horizon, generate predictions without refitting, estimate calibration quantities under an explicit model, and report uncertainty together with graphical assessments. A slope below one indicates predictions that are generally too extreme, whereas a slope above one indicates predictions that are insufficiently dispersed, provided that the slope is estimated on the original model's linear predictor. Recalibration, model updating, subgroup assessment, and clinical-utility analysis are treated as distinct tasks. The framework is intended to improve the traceability and interpretation of external evaluations; it does not convert statistical calibration into evidence of clinical usefulness.",
        east_asia="Times New Roman",
    )
    keywords = format_body(document.add_paragraph(), after=6, first_line=False)
    add_run(keywords, "Keywords  ", bold=True, east_asia="Times New Roman")
    add_run(
        keywords,
        "calibration slope; external validation; prediction model; reproducibility; reporting",
        east_asia="Times New Roman",
    )

    add_heading(document, "1  Introduction", east_asia="Times New Roman")
    add_body(
        document,
        "External evaluation asks whether a prediction model retains adequate performance when applied to data that were not used to develop it. Calibration is central to that question because many clinical decisions depend on agreement between predicted probabilities and observed risks, not only on the ordering of individuals. Nevertheless, evaluations often report one discrimination statistic and one calibration slope while omitting the precise model version, prediction horizon, analysis set, or scale on which recalibration was performed [1,2]. These omissions make apparently comparable estimates answer different questions.",
        east_asia="Times New Roman",
    )
    add_body(
        document,
        "The calibration slope is particularly vulnerable to reversed interpretation. When the slope is below one, differences in the original linear predictor are generally larger than supported by the evaluation data: high predictions tend to be too high and low predictions too low. When the slope exceeds one, the predictions vary too little and are concentrated toward the average risk [2,3]. These statements assume that the original predictor has been preserved and entered into a correctly specified recalibration model. A slope fitted after arbitrary rescaling, predictor refitting, or undocumented preprocessing does not retain the same interpretation.",
        east_asia="Times New Roman",
    )
    add_body(
        document,
        "A calibration slope is also not a complete model evaluation. Calibration-in-the-large addresses systematic overprediction or underprediction, flexible curves describe local agreement, discrimination describes ranking, and decision-analytic measures assess consequences under specified thresholds. A model can rank individuals well while producing poorly calibrated probabilities, and a recalibrated model can remain clinically unhelpful [1,4]. Each quantity therefore needs a declared estimand, an appropriate analysis set, and uncertainty that reflects the evaluation design.",
        east_asia="Times New Roman",
    )
    add_body(
        document,
        "The objective of this article is to define a practical evidence chain for interpreting calibration slopes in external model evaluation. The framework links the frozen prediction specification, target setting, evaluation dataset, fitted performance objects, machine-readable results, figures, and manuscript statements. It is designed as a neutral author-manuscript framework rather than a journal-layout template and is applicable to binary and time-to-event predictions with outcome-specific adaptations.",
        east_asia="Times New Roman",
    )

    add_heading(document, "2  Scope and estimands", east_asia="Times New Roman")
    add_heading(
        document,
        "2.1  Target population, setting, and prediction horizon",
        east_asia="Times New Roman",
        size=11,
        level=2,
    )
    add_body(
        document,
        "An external evaluation begins by defining where and for whom the model is intended to be used. Eligibility criteria, care setting, time zero, outcome definition, competing events, and prediction horizon must be stated before performance is examined. The analysis set should be derived from these definitions rather than from complete-case convenience. If several horizons or settings are clinically relevant, each constitutes a distinct evaluation target and should not be combined into one generic calibration statement.",
        east_asia="Times New Roman",
    )
    add_heading(
        document,
        "2.2  The prediction object",
        east_asia="Times New Roman",
        size=11,
        level=2,
    )
    add_body(
        document,
        "The evaluated object comprises more than a list of coefficients. It includes predictor definitions, measurement timing, transformations, coding rules, coefficients, intercept or baseline hazard, and any preprocessing learned in the development data. These components are frozen before evaluation. Predictions are generated without refitting the original coefficients. If a required transformation or baseline quantity is unavailable, the model is not fully transportable and the missing component should be reported rather than silently reconstructed from the evaluation sample.",
        east_asia="Times New Roman",
    )
    add_heading(
        document,
        "2.3  Calibration estimands",
        east_asia="Times New Roman",
        size=11,
        level=2,
    )
    add_body(
        document,
        "Calibration-in-the-large compares the average predicted risk with the observed risk and captures a systematic shift. The calibration slope evaluates whether the spread of the original predictions is appropriate. A flexible calibration curve examines local agreement across the prediction range. These estimands are related but not interchangeable. For survival outcomes, their definitions must specify the horizon and a method that accounts for censoring; for competing-risk outcomes, the predicted and observed quantities must refer to the same cumulative-incidence estimand.",
        east_asia="Times New Roman",
    )

    add_heading(document, "3  Evaluation framework", east_asia="Times New Roman")
    add_caption(
        document,
        "Table 1  Core components of a reproducible external calibration assessment",
        east_asia="Times New Roman",
    )
    add_table(
        document,
        ("Component", "Required specification", "Primary check", "Interpretive boundary"),
        (
            ("Prediction rule", "Version, coefficients, intercept or baseline hazard, transformations", "Reproduce predictions on test inputs", "Do not refit during evaluation"),
            ("Evaluation target", "Population, setting, time zero, outcome, horizon", "Match eligibility and endpoint definitions", "Do not merge distinct targets"),
            ("Analysis set", "Exclusions, missing data, censoring, competing events", "Account for all eligible records", "Complete cases may change the target"),
            ("Calibration", "Intercept, slope, curve, uncertainty", "Use the original prediction scale", "One summary cannot describe local agreement"),
            ("Clinical utility", "Decision, threshold range, consequences", "Compare decision strategies", "Calibration alone is insufficient"),
            ("Traceability", "Model object, result record, table, figure, text", "Reconcile values and direction", "Do not type corrections into outputs"),
        ),
        (25, 46, 42, 45),
        east_asia="Times New Roman",
        size=7.8,
    )

    add_heading(
        document,
        "3.1  Preserve the original prediction scale",
        east_asia="Times New Roman",
        size=11,
        level=2,
    )
    add_body(
        document,
        "For a logistic model, the original linear predictor is commonly entered as the sole covariate in a recalibration model. Estimating an intercept with the slope fixed at one assesses calibration-in-the-large; estimating both intercept and slope assesses the spread of predictions. The exact parameterization and link function should be reported. The same principle applies to survival models, but the baseline hazard, prediction horizon, and censoring-aware estimation must remain aligned with the original risk definition.",
        east_asia="Times New Roman",
    )
    add_body(
        document,
        "Preprocessing estimated in the development data, including centering, scaling, category cut points, and imputation models, is part of the frozen prediction rule. Re-estimating these quantities in the external dataset can improve apparent fit while changing the object under evaluation. Such a procedure may be legitimate model updating, but it must be labeled and analyzed separately from the untouched-model assessment.",
        east_asia="Times New Roman",
    )

    add_heading(
        document,
        "3.2  Define the analysis set and handle missing data",
        east_asia="Times New Roman",
        size=11,
        level=2,
    )
    add_body(
        document,
        "The report should account for the number of eligible individuals, exclusions before prediction, unavailable predictors, missing outcomes, and records included in each performance calculation. Missingness can alter both case mix and calibration. An imputation strategy should preserve the distinction between information available in the development model and information learned from the evaluation sample. Sensitivity analyses are appropriate when plausible missing-data assumptions lead to materially different performance conclusions.",
        east_asia="Times New Roman",
    )
    add_body(
        document,
        "For time-to-event outcomes, censoring and competing events require additional care. The observed outcome estimate used in a calibration plot must correspond to the same horizon and event definition as the predicted risk. Numbers at risk, follow-up distribution, and event counts provide essential context for uncertainty. A nominally large cohort may still yield imprecise calibration estimates if few events occur in clinically important parts of the prediction range.",
        east_asia="Times New Roman",
    )

    add_heading(
        document,
        "3.3  Quantify uncertainty and inspect local agreement",
        east_asia="Times New Roman",
        size=11,
        level=2,
    )
    add_body(
        document,
        "Calibration estimates should be accompanied by confidence intervals generated by a method appropriate to the sampling and outcome structure. A point estimate that differs from the ideal value can remain compatible with substantial sampling uncertainty. Conversely, a narrow interval around a non-ideal value can identify a reproducible mismatch even when the absolute numerical departure appears modest. Interpretation should therefore combine effect size, precision, and potential decision consequences.",
        east_asia="Times New Roman",
    )
    add_body(
        document,
        "A calibration curve should retain the probability scale and show the region in which evaluation data are available. Smoothing method, grouping, boundary behavior, and extrapolation can alter the visual impression. The curve should be displayed with uncertainty and the distribution of predicted risks. Decorative gradients, adequacy badges, or grouped observed-to-expected ratios cannot replace a curve when the scientific question concerns local agreement [2,3].",
        east_asia="Times New Roman",
    )

    add_heading(
        document,
        "3.4  Separate evaluation, recalibration, and updating",
        east_asia="Times New Roman",
        size=11,
        level=2,
    )
    add_body(
        document,
        "The untouched model should be evaluated first. Recalibration of the intercept or slope, revision of selected coefficients, addition of predictors, and full redevelopment represent progressively more extensive updating strategies. Their apparent performance should not be reported as if it described the original model. Updated models require internal validation within the updating dataset and, ideally, evaluation in further data because the updating step consumes information and can overfit.",
        east_asia="Times New Roman",
    )
    add_body(
        document,
        "Subgroup analyses also answer additional questions. A global calibration slope can conceal clinically important variation, but subgroup-specific estimates are often unstable and can encourage post hoc interpretation. Subgroups should be prespecified when possible, supported by adequate sample size and event counts, and accompanied by interaction or heterogeneity assessments that reflect the intended use rather than arbitrary significance thresholds.",
        east_asia="Times New Roman",
    )

    add_heading(document, "4  Reporting and interpretation", east_asia="Times New Roman")
    add_heading(
        document,
        "4.1  Minimum numerical and graphical reporting",
        east_asia="Times New Roman",
        size=11,
        level=2,
    )
    add_body(
        document,
        "A complete report identifies the model and version, target population, setting, outcome, prediction horizon, eligible sample, analysis set, number of events, missing-data handling, and prediction distribution. It reports calibration-in-the-large, calibration slope, confidence intervals, discrimination, and a calibration curve when these quantities are relevant. For survival outcomes, the amount of follow-up and censoring should be described. Each table and figure should state the evaluated model and analysis set rather than relying on a generic caption.",
        east_asia="Times New Roman",
    )
    add_body(
        document,
        "The manuscript should preserve the direction of each estimate. A slope below one may be described as predictions that are too extreme on average; a slope above one may be described as predictions that are insufficiently dispersed. Neither finding identifies which predictor coefficients are wrong, proves that development overfitting was the cause, or establishes that a particular updating strategy will improve decisions. Causal language is inappropriate unless supported by a design and analysis that address a causal question.",
        east_asia="Times New Roman",
    )

    add_heading(
        document,
        "4.2  Traceable assembly of results",
        east_asia="Times New Roman",
        size=11,
        level=2,
    )
    add_body(
        document,
        "Numerical results should be exported from the fitted evaluation objects to a stable machine-readable record. Tables, figures, and manuscript statements should read from that record or from verified derivatives. Corrections then return to the analysis that generated the estimate instead of being typed into separate outputs. At minimum, validation should reconcile sample sizes, event counts, model version, estimate direction, interval limits, rounding, and labels across the result record and all public-facing products.",
        east_asia="Times New Roman",
    )
    add_body(
        document,
        "A reproducible record does not guarantee a scientifically valid analysis. It makes the analysis inspectable. Design choices, outcome definitions, exclusions, missing-data assumptions, and updating decisions still require domain justification. Automated checks are most useful when they protect these declared contracts and stop on mismatches, while substantive interpretation remains tied to the clinical question and the evidence available in the evaluation setting.",
        east_asia="Times New Roman",
    )

    add_heading(document, "5  Discussion", east_asia="Times New Roman")
    add_body(
        document,
        "The proposed framework treats external calibration as an evidence chain rather than a single diagnostic number. Its central requirement is to preserve the identity of the prediction model while making the target setting and performance estimands explicit. This separation prevents an updated model from being mistaken for the original model and prevents a change in case mix, horizon, or outcome definition from being described as a universal property of model quality.",
        east_asia="Times New Roman",
    )
    add_body(
        document,
        "The framework also clarifies why calibration slope direction must be reported carefully. A slope below one indicates excessive spread of the original predictor relative to the evaluation data, whereas a slope above one indicates insufficient spread. The labels do not supply a mechanism. Differences in case mix, measurement, predictor effects, treatment patterns, or outcome incidence may contribute. Investigation of those mechanisms requires prespecified analyses and cannot be inferred from the slope alone.",
        east_asia="Times New Roman",
    )
    add_body(
        document,
        "Clinical utility remains a separate endpoint. Calibration affects the credibility of predicted probabilities, but a well-calibrated model may not improve decisions if it does not change management or if the consequences of errors outweigh benefits. Conversely, a model with imperfect overall calibration might retain value after transparent recalibration in a setting where its ranking information and decision consequences have been independently assessed [1,4]. Decision-curve or other utility analyses should therefore specify the action, threshold range, and relative consequences rather than using calibration as a proxy.",
        east_asia="Times New Roman",
    )
    add_body(
        document,
        "Several limitations should be recognized. The framework does not prescribe one estimator for every outcome or sampling design, and additional methods are needed for clustered data, complex surveys, dynamic predictions, recurrent events, and competing risks. It also does not provide universal sample-size thresholds. Precision depends on the prediction distribution, event frequency, censoring, and complexity of the calibration assessment. Finally, transparent reporting cannot remedy an evaluation dataset that is poorly aligned with the intended target population.",
        east_asia="Times New Roman",
    )
    add_body(
        document,
        "The article is methodological and does not present a new empirical dataset. Its value should therefore be judged by the clarity, applicability, and verifiability of the proposed assessment sequence. Future work could convert the framework into outcome-specific reporting templates and evaluate whether structured result records reduce direction, rounding, and version errors across analysis, figures, and manuscripts.",
        east_asia="Times New Roman",
    )

    add_heading(document, "6  Conclusion", east_asia="Times New Roman")
    add_body(
        document,
        "Calibration slopes can be interpreted reliably only when the original prediction rule, target setting, analysis set, recalibration model, uncertainty, and graphical evidence are all explicit. A slope below one generally indicates predictions that are too extreme, whereas a slope above one indicates predictions that are not dispersed enough. These findings do not establish clinical usefulness or identify a unique updating strategy. Linking fitted objects to machine-readable results and then to tables, figures, and manuscript statements provides a practical basis for reproducible external model evaluation.",
        east_asia="Times New Roman",
    )

    add_heading(document, "Declarations", east_asia="Times New Roman")
    add_body(
        document,
        "Ethics and data: This methodological example does not analyze participant-level data and therefore does not require participant consent or ethics review. Funding, affiliations, and conflicts of interest are not supplied because the demonstration has no named authors or institutions; they should not be fabricated to imitate a submitted article. Data availability: No empirical dataset was generated for this methods article. The document-generation source is available with the demonstration repository.",
        east_asia="Times New Roman",
        first_line=False,
    )

    add_heading(document, "References", east_asia="Times New Roman")
    references = (
        "[1] Steyerberg EW, Vickers AJ, Cook NR, et al. Assessing the performance of prediction models: a framework for traditional and novel measures. Epidemiology. 2010;21:128–138. doi:10.1097/EDE.0b013e3181c30fb2.",
        "[2] Van Calster B, McLernon DJ, van Smeden M, Wynants L, Steyerberg EW. Calibration: the Achilles heel of predictive analytics. BMC Medicine. 2019;17:230. doi:10.1186/s12916-019-1466-7.",
        "[3] Van Calster B, Nieboer D, Vergouwe Y, De Cock B, Pencina MJ, Steyerberg EW. A calibration hierarchy for risk models was defined: from utopia to empirical data. Journal of Clinical Epidemiology. 2016;74:167–176. doi:10.1016/j.jclinepi.2015.12.005.",
        "[4] Van Calster B, Vickers AJ. Calibration of risk prediction models: impact on decision-analytic performance. Medical Decision Making. 2015;35:162–169. doi:10.1177/0272989X14547233.",
    )
    for reference in references:
        add_reference(
            document,
            reference,
            east_asia="Times New Roman",
        )

    set_zoom(document)
    document.core_properties.title = (
        "A Reproducible Framework for Interpreting Calibration Slopes "
        "in External Model Evaluation"
    )
    document.core_properties.subject = "Complete single-column methodological author manuscript"
    document.save(EN_DOCX_PATH)


def render_with_word(docx_path):
    try:
        import fitz
        from win32com.client import DispatchEx
    except ImportError:
        return None

    pdf_path = docx_path.with_suffix(".pdf")
    png_path = docx_path.with_suffix(".png")
    word = DispatchEx("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    document = None
    try:
        document = word.Documents.Open(str(docx_path.resolve()), ReadOnly=True)
        document.Repaginate()
        document.ExportAsFixedFormat(str(pdf_path.resolve()), 17)
    finally:
        if document is not None:
            document.Close(False)
        word.Quit()

    with fitz.open(pdf_path) as pdf_document:
        page_count = pdf_document.page_count
        if page_count < 3:
            raise RuntimeError(
                f"{docx_path.name} 应为至少 3 页的完整稿，实际为 {page_count} 页"
            )
        page = pdf_document.load_page(0)
        pixmap = page.get_pixmap(matrix=fitz.Matrix(2.4, 2.4), alpha=False)
        pixmap.save(png_path)
    return png_path, page_count


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    required_paths = (
        RESULTS_PATH,
        COHORT_PATH,
        DIAGNOSTICS_PATH,
        FOREST_RESULTS_PATH,
        SURVIVAL_FIGURE_PATH,
        FOREST_FIGURE_PATH,
    )
    missing = [str(path) for path in required_paths if not path.exists()]
    if missing:
        raise FileNotFoundError("缺少论文生成所需的已验证结果：" + ", ".join(missing))

    result = read_csv_rows(RESULTS_PATH)[0]
    cohort_rows = read_csv_rows(COHORT_PATH)
    diagnostic_rows = read_csv_rows(DIAGNOSTICS_PATH)
    forest_rows = read_csv_rows(FOREST_RESULTS_PATH)

    build_chinese_manuscript(result, cohort_rows, diagnostic_rows, forest_rows)
    build_english_manuscript()
    for docx_path in (ZH_DOCX_PATH, EN_DOCX_PATH):
        rendered = render_with_word(docx_path)
        if rendered is None:
            print(f"已生成 {docx_path.name}；当前环境没有可用的 Word 渲染依赖")
        else:
            _, page_count = rendered
            print(f"已生成并渲染 {docx_path.name}，共 {page_count} 页")


if __name__ == "__main__":
    main()
