#!/usr/bin/env python3

import csv
import shutil
import subprocess
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


DEMO_DIR = Path(__file__).resolve().parent
OUTPUT_ROOT = DEMO_DIR / "output" / "document-skills"
FIGURE_RESULTS = DEMO_DIR / "output" / "publication-figures"

EN_FONT = "Times New Roman"
CN_BODY = "宋体"
CN_HEADING = "黑体"
BLACK = RGBColor(0, 0, 0)


def style_run(run, size=10.5, bold=False, italic=False, east_asia=CN_BODY):
    run.font.name = EN_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = BLACK
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), EN_FONT)
    rfonts.set(qn("w:hAnsi"), EN_FONT)
    rfonts.set(qn("w:eastAsia"), east_asia)
    return run


def add_segments(paragraph, segments, size=10.5, east_asia=CN_BODY):
    for segment in segments:
        if isinstance(segment, str):
            text, formatting = segment, {}
        else:
            text, formatting = segment
        style_run(
            paragraph.add_run(text),
            size=formatting.get("size", size),
            bold=formatting.get("bold", False),
            italic=formatting.get("italic", False),
            east_asia=formatting.get("east_asia", east_asia),
        )
    return paragraph


def configure_document(document):
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(22)
    section.right_margin = Mm(22)
    section.header_distance = Mm(8)
    section.footer_distance = Mm(8)

    normal = document.styles["Normal"]
    normal.font.name = EN_FONT
    normal.font.size = Pt(10.5)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), EN_FONT)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), EN_FONT)
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CN_BODY)

    zoom = document.settings._element.xpath("./w:zoom")
    if zoom:
        zoom[0].set(qn("w:percent"), "100")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = style_run(footer.add_run(), size=9)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, instruction, end))


def add_title(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.keep_with_next = True
    add_segments(
        paragraph,
        [(text, {"size": 16, "bold": True, "east_asia": CN_HEADING})],
    )
    return paragraph


def add_note(document, text):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    add_segments(paragraph, [(text, {"size": 9})], size=9)
    return paragraph


def add_heading(document, text, level=1):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    add_segments(
        paragraph,
        [
            (
                text,
                {
                    "size": 12 if level == 1 else 11,
                    "bold": True,
                    "east_asia": CN_HEADING,
                },
            )
        ],
    )
    return paragraph


def add_body(document, segments, first_line=True, size=10.5, after=5):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_after = Pt(after)
    if first_line:
        paragraph.paragraph_format.first_line_indent = Pt(size * 2)
    add_segments(paragraph, segments if isinstance(segments, list) else [segments], size=size)
    return paragraph


def add_numbered_item(document, label, segments, size=10.2):
    paragraph = document.add_paragraph(style="List Number")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.2
    paragraph.paragraph_format.space_after = Pt(3)
    add_segments(paragraph, [(label, {"bold": True})], size=size)
    add_segments(paragraph, segments if isinstance(segments, list) else [segments], size=size)
    return paragraph


def set_cell_border(cell, edge, value="nil", size=0, color="000000"):
    properties = cell._tc.get_or_add_tcPr()
    borders = properties.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        later_properties = {
            qn(f"w:{name}")
            for name in (
                "shd",
                "noWrap",
                "tcMar",
                "textDirection",
                "tcFitText",
                "vAlign",
                "hideMark",
            )
        }
        insertion_index = len(properties)
        for index, child in enumerate(properties):
            if child.tag in later_properties:
                insertion_index = index
                break
        properties.insert(insertion_index, borders)
    border = borders.find(qn(f"w:{edge}"))
    if border is None:
        border = OxmlElement(f"w:{edge}")
        borders.append(border)
    border.set(qn("w:val"), value)
    border.set(qn("w:sz"), str(size))
    border.set(qn("w:space"), "0")
    border.set(qn("w:color"), color)


def set_cell_segments(cell, segments, width_mm, bold=False, align="center", size=9):
    cell.text = ""
    cell.width = Mm(width_mm)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell.paragraphs[0]
    paragraph.alignment = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
    }[align]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.05
    normalized = []
    for segment in segments if isinstance(segments, list) else [segments]:
        if isinstance(segment, str):
            normalized.append((segment, {"bold": bold}))
        else:
            text, formatting = segment
            merged = dict(formatting)
            merged.setdefault("bold", bold)
            normalized.append((text, merged))
    add_segments(paragraph, normalized, size=size)


def add_three_line_table(document, headers, rows, widths_mm, aligns=None, size=9):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    aligns = aligns or ["center"] * len(headers)
    for index, (header, width) in enumerate(zip(headers, widths_mm)):
        table.columns[index].width = Mm(width)
        set_cell_segments(table.rows[0].cells[index], header, width, bold=True, align=aligns[index], size=size)
    for row in rows:
        cells = table.add_row().cells
        for index, (value, width) in enumerate(zip(row, widths_mm)):
            set_cell_segments(cells[index], value, width, align=aligns[index], size=size)
    for row in table.rows:
        for cell in row.cells:
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                set_cell_border(cell, edge)
    for cell in table.rows[0].cells:
        set_cell_border(cell, "top", value="single", size=8)
        set_cell_border(cell, "bottom", value="single", size=6)
    for cell in table.rows[-1].cells:
        set_cell_border(cell, "bottom", value="single", size=8)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def add_caption(document, segments, size=9):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.keep_with_next = True
    add_segments(paragraph, segments if isinstance(segments, list) else [segments], size=size)
    return paragraph


def read_single_row(path):
    with path.open(encoding="utf-8-sig", newline="") as handle:
        return next(csv.DictReader(handle))


def read_rows(path):
    with path.open(encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def save_document(document, path):
    path.parent.mkdir(parents=True, exist_ok=True)
    document.save(path)
    with zipfile.ZipFile(path) as package:
        bad_member = package.testzip()
        if bad_member:
            raise RuntimeError(f"DOCX ZIP 校验失败：{path} -> {bad_member}")


def build_study_design():
    output_dir = OUTPUT_ROOT / "epi-study-design"
    path = output_dir / "home-bp-monitoring-protocol-sap.docx"
    document = Document()
    configure_document(document)
    add_title(document, "家庭血压监测与收缩压控制的前瞻性队列研究方案")
    add_note(document, "演示性 PROTOCOL 与 SAP；所有设计参数须在真实实施前由研究责任人确认")

    add_heading(document, "一、研究方案（PROTOCOL）")
    add_body(
        document,
        "本演示方案拟评估社区高血压人群中家庭血压监测频率与 12 个月收缩压控制的关联。目标人群为 18～79 岁、已由医疗机构诊断高血压并计划在同一社区卫生服务机构随访的成年人。时间零点定义为完成基线诊室血压测量并确认纳入资格的日期。暴露为基线后前 4 周的家庭血压监测频率，比较每周至少 4 天与每周不足 1 天两组；该分组仅用于展示可执行定义，真实研究需在查看结局前结合设备记录完整性和临床意义确认。",
    )
    add_body(
        document,
        "主要结局为时间零点后 12 个月诊室收缩压相对基线的变化值，单位为 mmHg；次要结局包括 6 个月收缩压变化、12 个月舒张压变化和随访期间降压药物方案调整。主要 estimand 为目标人群中高频监测组与低频监测组的 12 个月收缩压平均变化差，汇总量为调整后的组间均值差。死亡、迁出和无法完成 12 个月测量作为中间事件分别记录，主分析不把其结果静默填补为无变化。观察性设计只能支持关联性解释，不能把组间差异写成家庭监测的因果治疗效果。",
    )

    add_caption(document, "表 1　研究问题与关键定义")
    add_three_line_table(
        document,
        ["项目", "操作性定义", "确认状态"],
        [
            ["目标人群", "18～79 岁、已诊断高血压、计划连续随访", "演示性预设"],
            ["时间零点", "基线诊室血压测量并确认纳入资格之日", "待责任人确认"],
            ["暴露比较", "每周至少 4 天与每周不足 1 天家庭测量", "待设备数据核验"],
            ["主要结局", "12 个月诊室收缩压相对基线变化（mmHg）", "待测量方案确认"],
            ["主要效应量", "调整后组间均值差及 95% 置信区间", "预设"],
        ],
        [31, 105, 30],
        aligns=["left", "left", "center"],
    )

    add_heading(document, "二、偏倚控制与可行性")
    add_body(
        document,
        "选择偏倚通过记录连续筛查人数、排除原因和未参加者特征进行评估；信息偏倚通过统一诊室测量流程、保留设备时间戳并在结局测量时不向测量人员展示分组信息进行控制。混杂因素预设包括基线收缩压、年龄、性别、糖尿病、慢性肾病、药物类别、用药依从性和就诊频率。药物调整既可能反映病情，也可能受家庭测量影响，因此主分析需明确其作为暴露后变量的处理，不能在未定义 estimand 时直接作为普通协变量调整。样本量或精度依据尚无经核验的先验标准差、组间差异和失访率，本演示稿将其标记为待确认，不以猜测数字补齐。",
    )

    document.add_page_break()
    add_heading(document, "三、统计分析计划（SAP）")
    add_body(
        document,
        "主要分析人群包括满足纳入条件、具有基线暴露判定且至少有一次结局随访的研究对象；是否采用全分析集及其与完全病例分析的关系须在实施前冻结。首先按暴露组描述基线特征和样本流，连续变量根据分布报告均数与标准差或中位数与四分位数，分类变量报告人数与构成比。组间描述不以单因素检验结果决定协变量纳入；若报告传统的独立样本 t 检验，统计符号按学术规范排版，并同时报告效应量和区间估计。",
    )
    add_body(
        document,
        [
            "主要模型拟采用线性回归估计 12 个月收缩压变化的调整后均值差，预设调整基线收缩压、年龄、性别、糖尿病、慢性肾病、药物类别、用药依从性和就诊频率。模型报告回归系数、95% 置信区间和必要的 ",
            ("P", {"bold": True, "italic": True}),
            " 值，检查残差分布、异方差和高影响观测。若结局分布或随访结构使线性模型不适用，任何替代模型均作为方案修订记录，不根据显著性选择。",
        ],
    )
    add_caption(document, "表 2　预设分析与边界")
    add_three_line_table(
        document,
        ["分析层级", "方法", "预设输出", "边界"],
        [
            ["主要分析", "多变量线性回归", "均值差、95% CI、P 值", "不作因果解释"],
            ["缺失数据", "描述模式；多重插补作为候选", "缺失比例、插补诊断", "机制假设待确认"],
            ["敏感性分析", "完全病例与不同暴露阈值", "方向与区间比较", "不替代主分析"],
            ["探索性分析", "年龄或合并症分层", "交互项与区间", "不据层内显著性下结论"],
        ],
        [27, 48, 46, 45],
        aligns=["left", "left", "left", "left"],
    )
    add_body(
        document,
        "实施前仍需由研究责任人确认家庭测量设备标准、有效测量日定义、主要暴露阈值、药物调整的 estimand 处理、12 个月测量时间窗、失访处置以及精度依据。上述事项会改变分析集或主要估计，未确认前不得进入主要分析，也不得把本演示稿视为已经冻结的正式方案。",
    )
    save_document(document, path)
    return path


def build_report():
    output_dir = OUTPUT_ROOT / "report-writing"
    path = output_dir / "fixed-cohort-survival-report.docx"
    result = read_single_row(FIGURE_RESULTS / "survival-demo-results.csv")
    forest = read_rows(FIGURE_RESULTS / "cox-forest-results.csv")
    n = int(result["n"])
    events = int(result["events"])
    hr = float(result["hazard_ratio"])
    low = float(result["ci_lower"])
    high = float(result["ci_upper"])
    p_value = float(result["p_value"])

    document = Document()
    configure_document(document)
    add_title(document, "固定模拟队列生存分析报告")
    add_note(document, "结果来自仓库固定模拟数据与实际 Cox 模型，仅用于展示可复核报告流程")
    add_heading(document, "一、结论摘要")
    add_body(
        document,
        [
            f"本次分析纳入 {n:,} 名固定模拟研究对象并观察到 {events:,} 个结局事件。调整年龄、性别、疾病分期和标准化生物标志物后，强化方案组相对于常规方案组的结局发生风险比为 ",
            ("HR", {}),
            f" {hr:.2f}（95% ",
            ("CI", {}),
            f" {low:.2f}～{high:.2f}，",
            ("P", {"bold": True, "italic": True}),
            " < 0.001）。在该固定模拟数据和预设模型下，强化方案与较低的结局发生风险相关；该结果不构成真实临床证据，也不能被解释为已经证实的治疗因果效应。",
        ],
    )
    add_heading(document, "二、数据与方法")
    add_body(
        document,
        "分析使用仓库内由固定随机种子生成的模拟队列，最长随访 36 个月。主要结局为首次模拟事件，未发生事件者在随访结束时行政删失。主要模型为 Cox 比例风险模型，治疗方案以常规方案为参照，并同时调整年龄（每增加 10 岁）、性别、疾病分期和标准化生物标志物。所有显示数字均由同一分析脚本输出，报告没有重新计算或手工改写关键结果。",
    )
    add_caption(document, "表 1　主要模型结果")
    result_rows = []
    for row in forest[:3]:
        p_num = float(row["p_value"])
        p_display = "< 0.001" if p_num < 0.001 else f"= {p_num:.3f}"
        result_rows.append(
            [
                row["label"],
                f"{float(row['estimate']):.2f}",
                f"{float(row['conf_low']):.2f}～{float(row['conf_high']):.2f}",
                [("P", {"bold": True, "italic": True}), f" {p_display}"],
            ]
        )
    add_three_line_table(
        document,
        ["模型项", "HR", "95% CI", "P 值"],
        result_rows,
        [79, 24, 38, 25],
        aligns=["left", "center", "center", "center"],
    )
    add_body(
        document,
        "表中强化方案、年龄和男性三项结果分别呈现效应方向、估计值和不确定性。其余疾病分期及生物标志物模型项保留在完整森林图和机器可读结果中，正文不逐格重复。模型项的参照水平、计量单位和 95% 置信区间与分析输出一致；报告没有在图内额外烧录解释性段落或教材式效应尺度说明。",
    )
    add_heading(document, "三、解释边界")
    add_body(
        document,
        "模拟队列中的治疗分配和结局风险由预设数据生成机制共同决定，数值只用于检查从分析脚本、结果文件到报告表述的可追溯性。风险比小于 1 表示在所拟合模型和参照设定下观察到较低的瞬时结局风险，不等于绝对风险降低，也不能替代真实研究中的混杂控制、模型诊断、敏感性分析和临床意义评价。若任何结果数字发生变化，应回到生成模型的脚本并同步更新结果文件、表格、图件和本报告。",
    )

    document.add_page_break()
    add_heading(document, "四、完整结果与复核位置")
    add_body(
        document,
        "完整结果包括调整后无事件生存曲线、风险集、全部模型项的森林图和模型诊断。生存曲线用于显示随访期间的估计概率与不确定性，森林图用于呈现治疗方案及协变量的风险比和 95% 置信区间。两个图件分别回答时间到事件分布和多变量模型结构，未拼成含有重复指标的看板。报告读者可从固定模拟数据重新运行分析脚本，并对照机器可读结果检查样本量、事件数、效应方向和显示精度。",
    )
    add_caption(document, "表 2　结果追溯位置")
    add_three_line_table(
        document,
        ["内容", "仓库位置", "核对重点"],
        [
            ["固定模拟数据", "docs/demo/survival-demo-data.csv", "样本与变量定义"],
            ["分析与出图脚本", "docs/demo/generate_survival_demo.R", "模型、参照与导出"],
            ["主要结果", "survival-demo-results.csv", "n、事件数、HR、CI、P"],
            ["完整模型项", "cox-forest-results.csv", "方向、区间与标签"],
        ],
        [38, 77, 51],
        aligns=["left", "left", "left"],
        size=8.5,
    )
    add_body(
        document,
        "复核时应同时检查脚本标准输出和标准错误，不能以退出码或日志末尾替代 warning、缺失值、样本量变化和模型异常扫描。若发现样本量、事件数、方向或区间与当前报告不一致，应停止外发并定位最早产生差异的文件；在无法确定哪个版本正确时，不按修改时间或显著性选择结果。",
    )
    save_document(document, path)
    return path


def build_reproducibility_memo():
    output_dir = OUTPUT_ROOT / "report-writing"
    path = output_dir / "fixed-cohort-reproducibility-memo.docx"
    data_path = DEMO_DIR / "survival-demo-data.csv"
    result_path = FIGURE_RESULTS / "survival-demo-results.csv"
    forest_path = FIGURE_RESULTS / "cox-forest-results.csv"
    script_path = DEMO_DIR / "generate_survival_demo.R"
    survival_figure = FIGURE_RESULTS / "adjusted-survival.png"
    forest_figure = FIGURE_RESULTS / "cox-forest.png"

    data_rows = read_rows(data_path)
    result = read_single_row(result_path)
    forest_rows = read_rows(forest_path)
    observed_n = len(data_rows)
    observed_events = sum(int(row["event"]) for row in data_rows)
    expected_n = int(result["n"])
    expected_events = int(result["events"])
    checks = [
        ["固定模拟数据", str(observed_n), str(expected_n), "通过" if observed_n == expected_n else "不通过"],
        ["结局事件数", str(observed_events), str(expected_events), "通过" if observed_events == expected_events else "不通过"],
        ["完整模型项", str(len(forest_rows)), "6", "通过" if len(forest_rows) == 6 else "不通过"],
        ["分析脚本", "文件存在" if script_path.is_file() else "缺失", "文件存在", "通过" if script_path.is_file() else "不通过"],
        ["调整后生存曲线", "文件存在" if survival_figure.is_file() else "缺失", "文件存在", "通过" if survival_figure.is_file() else "不通过"],
        ["森林图", "文件存在" if forest_figure.is_file() else "缺失", "文件存在", "通过" if forest_figure.is_file() else "不通过"],
    ]
    if any(row[-1] != "通过" for row in checks):
        raise RuntimeError("复现核查备忘录发现未通过项目，停止生成正式展示稿")

    document = Document()
    configure_document(document)
    add_title(document, "固定模拟队列成果复现核查备忘录")
    add_note(document, "核查对象为仓库公开演示文件，不涉及真实患者资料")
    add_heading(document, "一、核查结论")
    add_body(
        document,
        f"固定模拟数据共包含 {observed_n:,} 条记录和 {observed_events:,} 个结局事件，与机器可读结果文件中的样本量和事件数一致。完整森林图结果包含 {len(forest_rows)} 个预设模型项，分析脚本、调整后生存曲线和森林图均位于预期路径。当前核查支持 README 演示文件之间的数量与文件存在性一致，但不把文件存在视为模型科学有效性的充分证据，也不替代对完整标准输出、标准错误、warning、缺失值和模型诊断的审查。",
    )
    add_heading(document, "二、核查范围与方法")
    add_body(
        document,
        "本次核查从固定模拟 CSV 逐行计数样本和 event 字段，再与 survival-demo-results.csv 的 n 与 events 字段比较；同时读取 cox-forest-results.csv 的模型项数量，并检查分析脚本与两张统计图是否存在。核查没有重新拟合 Cox 模型，也没有改变任何数据、结果或图件。若数量或文件状态不一致，生成脚本会停止，不输出显示为“通过”的备忘录。",
    )
    add_caption(document, "表 1　复现核查结果")
    add_three_line_table(
        document,
        ["核查项目", "实际值", "预期值", "状态"],
        checks,
        [55, 42, 42, 27],
        aligns=["left", "center", "center", "center"],
    )
    add_heading(document, "三、解释与后续动作")
    add_body(
        document,
        "全部项目通过说明当前展示数据、结果摘要和文件路径可以相互对应。它不能证明没有数据错误、模型设定正确或结论可用于真实医学决策。正式分析仍需从项目根运行既定总脚本，扫描完整输出并检查模型假设、异常记录和实际使用这些数字的文件；任何关键数字变化都应回到生成结果的脚本，再同步更新结果文件、表图、论文、报告和演示文稿。",
    )
    save_document(document, path)
    return path


def build_peer_review():
    output_dir = OUTPUT_ROOT / "manuscript-peer-review"
    path = output_dir / "cohort-manuscript-review-report.docx"
    document = Document()
    configure_document(document)
    add_title(document, "观察性队列稿件同行评审报告")
    add_note(document, "公开演示情景；审查对象为虚构稿件片段，不对应真实作者、机构或未公开研究")
    add_heading(document, "一、审查依据与范围")
    add_body(
        document,
        "本报告审查一份用于演示的观察性队列稿件，材料包括摘要、方法、结果正文、表 1 和图 1。稿件描述家庭血压监测频率与 12 个月收缩压变化的关联，但未提供研究方案、注册记录、原始数据、分析代码或补充材料。报告依据稿件内可见信息核对研究问题、时间零点、样本流、变量定义、统计方法、结果与论断；STROBE 仅用于识别报告缺项，不作为研究质量评分。由于未取得原始数据和代码，本报告不能声称数字真实、分析已复现或数据中不存在错误。",
    )
    add_heading(document, "二、总体评价")
    add_body(
        document,
        "稿件的问题具有明确的公共卫生意义，暴露、主要结局和随访时长在摘要中可以识别，表 1 也报告了主要组间估计和 95% 置信区间。当前最重要的问题不是语言风格，而是方法中没有清楚定义时间零点和暴露判定窗口，缺失数据处理无法与样本流对应，摘要结论又把观察性关联写成家庭监测改善血压控制的因果效果。这些问题会影响偏倚判断和主要结论的可解释性；作者需要先补充报告、核对分析对象并降低论断强度，之后才能评价现有估计是否足以支持稿件主张。",
    )
    add_heading(document, "三、主要意见")
    add_heading(document, "M1　时间零点与暴露判定窗口未定义", level=2)
    add_body(
        document,
        "类型为 reporting gap。位置为方法“研究对象与暴露”第一段。稿件写明按家庭测量频率分组，却未说明从哪一天开始累计测量、暴露窗口是否先于结局随访，也未说明在暴露窗口内发生药物调整或失访时如何处理。读者因此无法判断暴露与结局的时间顺序，也无法排除将必须存活并完成测量的时间错误归入暴露组。请给出统一时间零点、暴露判定窗口、纳入资格和随访起点，并用样本流说明每一步人数；若暴露随时间更新，应说明相应的时间变化分析。核验状态为需作者澄清。",
    )
    add_heading(document, "M2　缺失数据与模型样本量不能对应", level=2)
    add_body(
        document,
        "类型为 inconsistency。位置为结果第二段和表 1。正文称主要模型纳入全部 842 名研究对象，表 1 的结局可用人数合计为 801 名，方法只写“缺失值予以排除”，未说明缺失发生在哪些变量、排除顺序或模型实际分母。该差异可能改变估计精度并影响缺失偏倚判断。请按原始、符合资格、暴露可判定、结局可用和主要模型分析集依次报告人数，同时说明各变量缺失量、完全病例或插补方法及其假设；若不同模型使用不同样本，应逐项报告分析人数。核验状态为稿件内已定位、未取得原始数据。",
    )
    add_heading(document, "M3　摘要结论超过观察性证据", level=2)
    add_body(
        document,
        "类型为 interpretation。位置为摘要结论和讨论末段。稿件使用“家庭监测可改善血压控制并应推广”的确定性措辞，但设计为非随机观察性队列，现有材料没有说明足以识别因果效应的策略，也没有报告绝对差异、伤害、成本或实施可行性。请将主要结论改为家庭监测频率与收缩压变化之间的调整后关联，明确残余混杂和选择偏倚，并把推广建议限定为需要进一步因果评价的研究问题。若作者坚持因果主张，需要说明 estimand、识别假设、混杂控制依据和相应敏感性分析。核验状态为稿件内已核对。",
    )

    document.add_page_break()
    add_heading(document, "四、次要意见")
    add_numbered_item(document, "m1　表 1：", "请在表题或表注中定义调整模型、参照组、分析人数和 95% 置信区间，保持所有单元格白底并使用三线表。")
    add_numbered_item(
        document,
        "m2　统计符号：",
        [
            "英文与数字使用 Times New Roman；",
            ("P", {"bold": True, "italic": True}),
            " 使用粗斜体，其余拉丁统计符号如 ",
            ("t", {"italic": True}),
            "、",
            ("F", {"italic": True}),
            "、",
            ("z", {"italic": True}),
            " 使用斜体，HR、CI 和数值保持正体。",
        ],
    )
    add_numbered_item(document, "m3　图 1：", "图内只保留坐标、单位、图例和必要数值，不重复摘要结论或添加教材式解释；完整方法和限制放入正文或图注。")

    add_heading(document, "五、覆盖矩阵")
    add_three_line_table(
        document,
        ["维度", "当前判断", "问题编号", "未核验边界"],
        [
            ["研究问题与贡献", "问题可识别", "M3", "未核验外部创新性"],
            ["数据与稿内一致性", "存在分母冲突", "M2", "无原始数据"],
            ["研究设计与偏倚", "时间顺序不清", "M1", "无方案与注册"],
            ["统计方法与不确定性", "目前无法完整判断", "M1、M2", "无代码与诊断"],
            ["结果、讨论与结论", "论断强度过高", "M3", "未复算结果"],
            ["语言、结构与表图", "需局部规范化", "m1～m3", "仅审所给材料"],
        ],
        [39, 43, 31, 53],
        aligns=["left", "left", "center", "left"],
        size=8.5,
    )
    add_heading(document, "六、未核验事项")
    add_body(
        document,
        "本报告未核验作者身份、研究机构、伦理批准、原始数据、分析代码、参考文献真实性、期刊最新表单或稿件外的研究实施过程。报告没有给出录用、修改或拒稿建议，因为演示情景未提供目标期刊的决定类别和评价标准。若后续取得方案、数据或代码，应分别核对预设终点、样本流、主要模型、缺失处理和表图数字；只有实际运行授权代码后，才可把相应项目标记为已复算。",
    )
    save_document(document, path)
    return path


def workflow_text():
    return """工作流问题交接报告
报告时间：2026-08-04
报告范围：EpiAgentKit README 科研内容图展示的制作与验收

一、结论摘要
本次问题的共同原因是正文内容图在制作阶段被误当作氛围插图，提示词先描述画风和场景，没有先锁定需要呈现的对象、字段、关系与流程状态；验收又只检查整体观感，没有执行内容删除测试和逐字逐边核对。结果是图片可以作为网页背景，但不能说明 skill 的真实输入、处理和输出。

二、证据边界
用户确认：除网页装饰或明确特殊场景外，不接受没有具体任务内容的概念插图。
直接观察：当前修订后的证据核验图位于 docs/showcase/illustrations/evidence-research.png，包含研究问题、来源类型、证据矩阵和采用/补充检索/排除状态。
可复核推断：最早失效发生在图片类型判定和提示词内容规划，而不是导出分辨率或 README 嵌入阶段。
待核验：第一次空泛候选未作为正式仓库成果保留，后续维护者只能依据会话纠正和当前成果核对差异。

三、问题记录
问题 WF-001
涉及的工作项、产物与阶段：README 中 evidence-research、consulting-delivery、epiagentkit-maintenance 和 academic-humanizer 的 image2 展示图；发生在图前规划与第一轮验收。
适用条件、不适用范围与合法例外：适用于解释 skill、研究流程、架构、机制和成果路径的正文内容图；不适用于纯网页装饰、封面氛围图或用户明确要求的特殊场景。
用户确认的正确结果：内容图应直接显示具体对象、字段、关系、流程状态或真实截图依据，脱离相邻说明仍能识别图中任务。
实际表现：初始候选以抽象卡片、设备和背景氛围表达 evidence-research，缺少可核验的研究问题、来源身份、偏倚、适用性和决策关系。
证据位置与证据类型：用户会话纠正属于用户确认；docs/showcase/illustrations/evidence-research.png 属于修订后直接观察。
当时适用的规则或 skill：research-visuals 应先区分正文内容图、真实截图和氛围插图，再按 scenario-playbook.md 规划最小充分文字和关系。
最早失效环节：图片被误判为 README 氛围插图，提示词先锁定审美方向而未先列出内容对象和来源对应关系。
现有流程未能阻止的原因：第一轮验收没有逐项比较用户要求、图中对象、文字和箭头，也没有把“是否能说明真实任务”设为完成条件。
影响：图片无法帮助读者理解 skill 的输入、判断过程和可交付结果，README 展示与实际能力脱节。
置信度与待核验内容：共同原因置信度高；初始候选文件未入仓，无法进行像素级前后比较。

四、共同原因
内容类型判定、提示词内容规划和验收标准之间没有相互对应。只强调审美、纯色和信息密度，不能代替对具体对象、字段、关系和来源的核对。

五、交给 EpiAgentKit 核验的候选调整
目标行为：正文内容图先建立最小充分内容清单，再设计视觉系统；只有网页装饰或明确特殊场景允许空泛氛围图。
执行主体、动作与完成证据：research-visuals 在生成前列出对象、字段、关系和逐字文字，imagegen 生成后逐项核对；完成证据为两份内容要求不同且分别合格的可打开成果及 review/INDEX.md。
触发条件与不触发条件：README 工作流、架构、skill 示例和科研机制图触发；真实统计图转 publication-figures，真实界面使用截图，科研原始图像不得生成式重绘。
最小机制与可能位置：全局规则保留内容优先边界；research-visuals/SKILL.md 负责分流；scenario-playbook.md 和 prompt-recipes.md 负责条件细则与提示词模板；维护测试检查 README 示例不再链接空泛概念图。
必须保留：封面、网页装饰和用户明确选择的特殊氛围场景仍可使用无字或少字插图；不能把本次内容规则扩大为所有图片必须密集排版。
旧场景与新场景：旧场景为网页 hero 氛围图，仍应通过；新场景为 evidence-research 方法依据核验图，必须显示研究问题、来源、证据矩阵和决策状态。
证据限制：候选位置需在完整仓库中核对，workflow.txt 不直接决定最终修改文件。

六、不应升级为通用规则的事项
本次使用的海军蓝、青绿和白色配色只属于 README 当前视觉系统，不能要求所有科研图采用同一配色。四个 skill 的横向构图也不是其它领域图件的固定模板。

七、尚缺证据
无初始空泛候选的仓库内文件；无第三方使用者在窄屏和不同 README 渲染器中的可读性反馈。
"""


def build_workflow_retrospective():
    output_dir = OUTPUT_ROOT / "workflow-retrospective"
    txt_path = output_dir / "workflow.txt"
    docx_path = output_dir / "workflow-retrospective-display.docx"
    output_dir.mkdir(parents=True, exist_ok=True)
    txt_path.write_text(workflow_text(), encoding="utf-8")

    document = Document()
    configure_document(document)
    add_title(document, "科研内容图空泛问题交接报告")
    add_note(document, "workflow.txt 的 Word 展示副本；正式交接内容以同目录 UTF-8 文本为准")
    add_heading(document, "一、结论摘要")
    add_body(
        document,
        "本次问题的共同原因是 README 正文内容图在制作阶段被误当作氛围插图。提示词先描述画风、设备和场景，没有先锁定需要呈现的研究对象、字段、关系与流程状态；第一轮验收又只检查整体观感，没有执行内容删除测试和逐字逐边核对。结果是图片可以充当网页背景，却不能说明 evidence-research 等 skill 的真实输入、判断过程和输出。",
    )
    add_heading(document, "二、证据边界")
    add_three_line_table(
        document,
        ["证据类型", "内容", "位置或限制"],
        [
            ["用户确认", "除网页装饰或明确特殊场景外，不接受空泛概念插图", "当前会话纠正"],
            ["直接观察", "修订图包含研究问题、来源、证据矩阵和决策状态", "docs/showcase/illustrations/evidence-research.png"],
            ["可复核推断", "最早失效位于图片类型判定与内容规划", "由当前图和规则对照得出"],
            ["待核验", "初始空泛候选未作为正式仓库成果保留", "不能进行像素级前后比较"],
        ],
        [31, 79, 56],
        aligns=["left", "left", "left"],
        size=8.3,
    )
    add_heading(document, "三、问题 WF-001")
    add_body(
        document,
        "涉及的工作项为 evidence-research、consulting-delivery、epiagentkit-maintenance 和 academic-humanizer 的 README image2 展示图，问题发生在图前规划与第一轮验收。正确结果是内容图直接显示具体对象、字段、关系、流程状态或真实截图依据，脱离相邻说明仍能识别图中任务。初始表现则以抽象卡片、设备和背景氛围表达 evidence-research，缺少研究问题、来源身份、偏倚、适用性和决策关系。",
    )
    add_body(
        document,
        "最早失效环节是把 README skill 示例误判为氛围图，导致提示词先锁定审美方向而未先列出最小充分内容。现有流程没有阻止问题，是因为验收没有逐项比较用户要求、图中对象、文字和箭头，也没有把“是否能说明真实任务”作为完成条件。该问题使展示图无法帮助读者理解 skill 的输入、判断和交付结果，README 视觉效果与实际能力发生脱节。共同原因的置信度高，但初始候选未入仓，因此无法做像素级比较。",
    )

    document.add_page_break()
    add_heading(document, "四、候选调整")
    add_body(
        document,
        "候选目标行为是让正文内容图先建立最小充分内容清单，再设计视觉系统。research-visuals 应在生成前列出对象、字段、关系和逐字文字，imagegen 返回后逐项核对；修改 skill 时另按维护流程生成两份内容要求不同且分别合格的可打开成果，并用 review/INDEX.md 说明验收重点。真实统计图继续转 publication-figures，真实界面使用截图，科研原始图像不得生成式重绘。",
    )
    add_heading(document, "五、必须保留与合法例外")
    add_body(
        document,
        "封面、网页装饰和用户明确选择的特殊氛围场景仍可使用无字或少字插图，本次调整不能被扩大为所有图片都要密集排版。海军蓝、青绿和白色只属于 README 当前视觉系统，不能要求医学机制图、机器学习架构图和学术汇报都使用同一配色或横向四栏构图。完成维护时应以旧的网页 hero 场景继续通过、新的证据核验内容图能够显示研究问题、来源、证据矩阵和决策状态作为一组新旧回归用例。",
    )
    add_heading(document, "六、尚缺证据")
    add_body(
        document,
        "当前缺少初始空泛候选的仓库内文件，也缺少第三方使用者在窄屏和不同 README 渲染器中的可读性反馈。后续维护者应在完整仓库中核对全局规则、research-visuals 核心流程、scenario-playbook.md、prompt-recipes.md、README 调用位置和相应测试，再决定最终修改位置；本交接报告只提供现场证据和候选机制，不预先指定必须修改或不得修改某个组件。",
    )
    save_document(document, docx_path)
    return txt_path, docx_path


def render_docx(docx_path, preview_page=1):
    pandoc = shutil.which("pandoc")
    xelatex = shutil.which("xelatex")
    pdftoppm = shutil.which("pdftoppm")
    if not (pandoc and xelatex and pdftoppm):
        missing = [name for name, path in (("pandoc", pandoc), ("xelatex", xelatex), ("pdftoppm", pdftoppm)) if not path]
        raise RuntimeError("缺少文档渲染依赖：" + ", ".join(missing))
    pdf_path = docx_path.with_suffix(".pdf")
    subprocess.run(
        [
            pandoc,
            str(docx_path),
            "--output",
            str(pdf_path),
            "--pdf-engine",
            xelatex,
            "-V",
            "papersize=a4",
            "-V",
            "geometry:top=20mm,bottom=18mm,left=22mm,right=22mm",
            "-V",
            f"mainfont={EN_FONT}",
            "-V",
            "CJKmainfont=SimSun",
            "-V",
            "fontsize=10pt",
        ],
        check=True,
        cwd=docx_path.parent,
    )
    preview_path = docx_path.with_suffix(".png")
    prefix = preview_path.with_suffix("")
    subprocess.run(
        [
            pdftoppm,
            "-png",
            "-r",
            "200",
            "-f",
            str(preview_page),
            "-l",
            str(preview_page),
            "-singlefile",
            str(pdf_path),
            str(prefix),
        ],
        check=True,
    )
    if not preview_path.exists():
        raise RuntimeError(f"未生成页面预览：{preview_path}")
    return pdf_path, preview_path


def main():
    outputs = [
        build_study_design(),
        build_report(),
        build_reproducibility_memo(),
        build_peer_review(),
    ]
    workflow_txt, workflow_docx = build_workflow_retrospective()
    outputs.append(workflow_docx)
    print(f"已生成 {workflow_txt.relative_to(DEMO_DIR.parent.parent)}")
    for docx_path in outputs:
        pdf_path, preview_path = render_docx(docx_path)
        print(f"已生成 {docx_path.relative_to(DEMO_DIR.parent.parent)}")
        print(f"已渲染 {pdf_path.relative_to(DEMO_DIR.parent.parent)}")
        print(f"已渲染 {preview_path.relative_to(DEMO_DIR.parent.parent)}")


if __name__ == "__main__":
    main()
