import importlib.util
import tempfile
import unittest
from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[2]
VALIDATOR_PATH = ROOT / "skills" / "graduate-opening-report" / "scripts" / "validate_proposal.py"
SPEC = importlib.util.spec_from_file_location("opening_report_validator", VALIDATOR_PATH)
VALIDATOR = importlib.util.module_from_spec(SPEC)
assert SPEC.loader is not None
SPEC.loader.exec_module(VALIDATOR)


class OpeningReportValidatorTests(unittest.TestCase):
    def make_showcase_docx(self, path):
        document = Document()
        section_terms = [alternatives[0] for alternatives in VALIDATOR.SECTION_GROUPS.values()]
        design_terms = [alternatives[0] for alternatives in VALIDATOR.DESIGN_GROUPS["observational"]["required"].values()]
        document.add_paragraph("结构完整性测试模板，含显式测试内容，不得用于学院归档。")
        document.add_paragraph(" ".join(section_terms + design_terms + ["观察性队列研究"]))
        target_fill = max(VALIDATOR.MIN_SHOWCASE_FILL_CHARS, VALIDATOR.MIN_SHOWCASE_TEXT_CHARS) + 1000
        fill_per_block = target_fill // VALIDATOR.MIN_SHOWCASE_FILL_BLOCKS + 1
        for index in range(VALIDATOR.MIN_SHOWCASE_FILL_BLOCKS):
            document.add_paragraph(
                f"第 {index + 1} 个模块应完成对应研究功能。"
                "【结构测试填充：以下 x 只检验篇幅和装配，不是正式研究内容】"
                + "x" * fill_per_block
            )
        document.add_paragraph("参考文献")
        for index in range(1, 9):
            document.add_paragraph(f"[{index}] 【引文待核验：第 {index} 类证据来源】")
        document.add_paragraph("附件与评议")
        for _ in range(VALIDATOR.MIN_DOCX_TABLES):
            table = document.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "字段"
            table.cell(0, 1).text = "内容"
            table.cell(1, 0).text = "测试"
            table.cell(1, 1).text = "【待补充：正式内容】"
        document.save(path)

    def test_showcase_passes_structural_mode_and_fails_archive_mode(self):
        with tempfile.TemporaryDirectory() as task_temp:
            path = Path(task_temp) / "complete-showcase.docx"
            self.make_showcase_docx(path)
            showcase = VALIDATOR.validate(path, mode="showcase")
            archive = VALIDATOR.validate(path, mode="archive")
        self.assertEqual(showcase["status"], "pass")
        self.assertEqual(showcase["showcase_fill_blocks"], VALIDATOR.MIN_SHOWCASE_FILL_BLOCKS)
        self.assertGreaterEqual(showcase["citation_pending_count"], 8)
        self.assertEqual(archive["status"], "fail")
        self.assertTrue(any("结构测试填充" in item for item in archive["errors"]))
        self.assertTrue(any("未解决占位符" in item for item in archive["errors"]))

    def test_short_outline_cannot_pass_showcase_or_archive(self):
        with tempfile.TemporaryDirectory() as task_temp:
            path = Path(task_temp) / "short-outline.docx"
            document = Document()
            document.add_paragraph("立题背景 研究目的 研究方法")
            document.save(path)
            showcase = VALIDATOR.validate(path, mode="showcase")
            archive = VALIDATOR.validate(path, mode="archive")
        self.assertEqual(showcase["status"], "fail")
        self.assertEqual(archive["status"], "fail")
        self.assertTrue(any("篇幅不足以证明完整报告深度" in item for item in showcase["errors"]))


if __name__ == "__main__":
    unittest.main()
