"""报告 docx 构建助手（report-writing skill 配套）。

用途：在用户没有指定其他模板时，把已经写好的报告内容排成中性 Word 文档，采用
中文宋体 / 英文 Times New Roman、三线表、表上图下题注、纯黑、
干净的居中加粗文首标题；页面保持白底，以字号、字重和间距建立层级。

用法（在生成脚本里 import）：
    from build_report import Report
    rep = Report()
    rep.title_lines(["一项 ... 研究", "24--36 周体重反弹补充分析报告"])  # 多行居中加粗，无副标题灰字
    rep.heading("一、分析背景与目的", level=1)
    rep.para("本补充分析用于评估 ...")                 # 完整段落
    rep.table_caption("表1 24--36周体重反弹分析样本量")
    rep.three_line_table(header=[...], rows=[...])      # 或 rep.table_from_xlsx(path, sheet)
    rep.three_line_table(header=["效应", "P 值"], rows=[["HR 0.74", "P < 0.001"]])
    rep.note("注：随机入组 N 按 ...")
    rep.figure(figure_paths["trajectory"], caption="图1 各组体重变化轨迹")  # 使用表图登记表或已经确认的输出路径
    rep.save("报告.docx", also_md=False)                # 仅在用户要求双格式时设 True

正文内容必须由调用方按 skill 强制要求写入（数据有源、完整段落、零编造），
本模块只负责"排版正确"，不负责"内容生成"。
"""
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CN_BODY = "宋体"           # 中文正文
CN_HEAD = "黑体"           # 中文标题（可换"微软雅黑"）
EN_FONT = "Times New Roman"  # 英文/数字
BLACK = (0, 0, 0)


def val(yaml_path, key, which="full"):
    """从 results.yaml 按固定结果名称读取已经格式化的结果文字。
    数字变化时先重新运行实际生成结果的脚本，再更新使用该结果的文件。
    用法：rep.para("S2 vs S1 差异为 " + val("results/results.yaml", "S2_vs_S1_diff") + "。")
    """
    import yaml
    with open(yaml_path, "r", encoding="utf-8") as f:
        doc = yaml.safe_load(f) or {}
    res = (doc.get("results") or {}).get(key)
    if res is None:
        raise KeyError(f"results.yaml 无键：{key}")
    display = res.get("display")
    if not isinstance(display, dict):
        rendered = res.get("rendered") or {}
        legacy_names = {
            "estimate": "est",
            "interval": "ci",
            "p_value": "p",
            "full": "full",
        }
        display = {name: rendered.get(old) for name, old in legacy_names.items()}
    s = display.get(which)
    if s is None:
        raise KeyError(f"结果 {key} 没有 display.{which}")
    return s


def setfont(run, cn=CN_BODY, en=EN_FONT, size=10.5, bold=False, italic=False, color=BLACK):
    """每个 run 都同时设置英文字体与中文 eastAsia 字体，避免使用非预期字体。"""
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = en
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), en)
    rfonts.set(qn("w:hAnsi"), en)
    rfonts.set(qn("w:eastAsia"), cn)
    run.font.color.rgb = RGBColor(*color)


def _set_cell_border(cell, **edges):
    """给单元格设指定边框（用于三线表：只在需要的行设顶/底线）。"""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        later_properties = {
            qn(f"w:{name}")
            for name in (
                "shd",
                "noWrap",
                "tcMar",
                "textDirection",
                "tcFitText",
                "vAlign",
                "hideMark",
            )
        }
        insertion_index = len(tcPr)
        for index, child in enumerate(tcPr):
            if child.tag in later_properties:
                insertion_index = index
                break
        tcPr.insert(insertion_index, borders)
    for edge in ("top", "left", "bottom", "right"):
        spec = edges.get(edge)
        tag = qn(f"w:{edge}")
        el = borders.find(tag)
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        if spec:  # spec = (sz_eighths_pt,)
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(spec[0]))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
        else:
            el.set(qn("w:val"), "none")


def _iter_explicit_runs(value, bold_default=False):
    items = value if isinstance(value, list) else ["" if value is None else str(value)]
    for item in items:
        if isinstance(item, tuple):
            text, opts = item
            opts = opts or {}
        else:
            text, opts = str(item), {}
        yield text, opts.get("bold", bold_default), opts.get("italic", False)


def _write_runs(paragraph, value, size, bold_default=False):
    """写普通值或显式 run 列表。"""
    for text, bold, italic in _iter_explicit_runs(value, bold_default):
        setfont(paragraph.add_run(text), size=size, bold=bold, italic=italic)


def _markdown_runs(value, bold_default=False):
    markdown = []
    for text, bold, italic in _iter_explicit_runs(value, bold_default):
        if bold and italic:
            markdown.append(f"***{text}***")
        elif italic:
            markdown.append(f"*{text}*")
        elif bold:
            markdown.append(f"**{text}**")
        else:
            markdown.append(text)
    return "".join(markdown)


class Report:
    def __init__(self, body_size=10.5):
        self.doc = Document()
        self.body_size = body_size
        self._md = []  # 同步累积 markdown
        zoom = self.doc.settings._element.xpath("./w:zoom")
        if zoom:
            zoom[0].set(qn("w:percent"), "100")
        # 默认正文样式字体（兜底；真正生效靠每个 run setfont）
        normal = self.doc.styles["Normal"]
        normal.font.name = EN_FONT
        normal.font.size = Pt(body_size)
        normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), CN_BODY)

    # ---- 文首标题 ----
    def title_lines(self, lines, size=16):
        """多行居中加粗文首标题，采用白底黑字。"""
        for ln in lines:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(4)
            setfont(p.add_run(ln), cn=CN_HEAD, size=size, bold=True)
        self.doc.add_paragraph()
        self._md.append("# " + " ".join(lines) + "\n")

    def meta(self, text):
        """可选的日期或版本，采用纯黑正常字号并居中。"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        setfont(p.add_run(text), size=self.body_size)
        self._md.append(f"*{text}*\n")

    # ---- 章节 ----
    def heading(self, text, level=1):
        size = {1: 15, 2: 13, 3: 12}.get(level, 12)
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
        setfont(p.add_run(text), cn=CN_HEAD, size=size, bold=True)
        self._md.append("\n" + "#" * (level + 1) + " " + text + "\n")

    # ---- 段落 ----
    def para(self, text, size=None):
        """普通完整段落：首行缩进 2 字符 + 1.5 倍行距（论文式正文）。"""
        sz = size or self.body_size
        p = self.doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Pt(2 * sz)   # 首行缩进 2 字符
        p.paragraph_format.space_after = Pt(4)
        setfont(p.add_run(text), size=sz)
        self._md.append(text + "\n")

    def para_runs(self, runs, size=None):
        """混合排版段落。"""
        sz = size or self.body_size
        p = self.doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.first_line_indent = Pt(2 * sz)   # 首行缩进 2 字符
        p.paragraph_format.space_after = Pt(4)
        md = []
        for text, opts in runs:
            opts = opts or {}
            setfont(p.add_run(text), size=sz,
                    italic=opts.get("italic", False), bold=opts.get("bold", False))
            if opts.get("bold") and opts.get("italic"):
                md.append(f"***{text}***")
            elif opts.get("italic"):
                md.append(f"*{text}*")
            elif opts.get("bold"):
                md.append(f"**{text}**")
            else:
                md.append(text)
        self._md.append("".join(md) + "\n")

    def summary_item(self, label, text):
        """当前章节内的并列提要段落：加粗标签 + 整句。"""
        p = self.doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(3)
        setfont(p.add_run(label + "："), size=self.body_size, bold=True)
        setfont(p.add_run(text), size=self.body_size)
        self._md.append(f"**{label}：**{text}\n")

    # ---- 表 ----
    def table_caption(self, text):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER   # 表题注居中（表上方）
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        setfont(p.add_run(text), size=self.body_size, bold=True)
        self._md.append("\n**" + text + "**\n")

    def note(self, text):
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        setfont(p.add_run(text), size=self.body_size - 1.5, italic=True)
        self._md.append(f"*{text}*\n")

    def three_line_table(self, header, rows, num_cols_right=None):
        """白底黑字三线表：仅设置顶线、表头下线和底线。
        num_cols_right: 右对齐的列索引集合（数字列）；默认除第 1 列外全右对齐。
        单元格可传普通值，或传由字符串和 ``(文字, {bold, italic})`` 组成的 run 列表。
        """
        ncol = len(header)
        if num_cols_right is None:
            num_cols_right = set(range(1, ncol))
        t = self.doc.add_table(rows=1, cols=ncol)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        big, small = 12, 4  # 边框粗细（1/8 pt）
        # 表头
        for j, htext in enumerate(header):
            cell = t.rows[0].cells[j]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _write_runs(p, htext, size=self.body_size, bold_default=True)
            _set_cell_border(cell, top=(big,), bottom=(small,))
        # 数据行
        for i, row in enumerate(rows):
            cells = t.add_row().cells
            last = (i == len(rows) - 1)
            for j, val in enumerate(row):
                cell = cells[j]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                p = cell.paragraphs[0]
                p.alignment = (WD_ALIGN_PARAGRAPH.RIGHT if j in num_cols_right
                               else WD_ALIGN_PARAGRAPH.LEFT)
                _write_runs(p, val, size=self.body_size)
                _set_cell_border(cell, bottom=(big,) if last else None)
        # markdown 镜像
        header_md = [_markdown_runs(value, bold_default=True) for value in header]
        self._md.append("| " + " | ".join(header_md) + " |")
        self._md.append("| " + " | ".join(["---"] * ncol) + " |")
        for row in rows:
            row_md = [_markdown_runs(value) for value in row]
            self._md.append("| " + " | ".join(row_md) + " |")
        self._md.append("")
        return t

    def table_from_xlsx(self, path, sheet=0, header_row=0, max_rows=None):
        """从 03_tables/ 的 xlsx 读一张表填三线表（自动取数，不手敲）。"""
        from openpyxl import load_workbook
        wb = load_workbook(path, data_only=True)
        ws = wb.worksheets[sheet] if isinstance(sheet, int) else wb[sheet]
        data = [[c if c is not None else "" for c in r]
                for r in ws.iter_rows(values_only=True)]
        data = [r for r in data if any(str(c).strip() for c in r)]  # 去空行
        header = data[header_row]
        rows = data[header_row + 1:]
        if max_rows:
            rows = rows[:max_rows]
        return self.three_line_table(header, rows)

    # ---- 图 ----
    def figure(self, path, caption=None, width_in=6.0):
        source = Path(path)
        embed = source
        if source.suffix.lower() == ".svg":
            embed = source.with_suffix(".png")
            if not embed.is_file():
                raise FileNotFoundError(
                    f"SVG 图解缺少可供当前文档嵌入的同名 PNG：{embed}。先按 research-visuals/svg-diagrams 生成预览图。"
                )
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(embed), width=Inches(width_in))
        if caption:
            c = self.doc.add_paragraph()
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.paragraph_format.space_before = Pt(2)
            setfont(c.add_run(caption), size=self.body_size, bold=True)
        self._md.append(f"\n![{caption or ''}]({source.as_posix()})")
        if caption:
            self._md.append(f"**{caption}**\n")

    # ---- 保存文件 ----
    def save(self, docx_path, also_md=False):
        self.doc.save(docx_path)
        out = [docx_path]
        if also_md:
            md_path = docx_path.rsplit(".", 1)[0] + ".md"
            with open(md_path, "w", encoding="utf-8") as f:
                f.write("\n".join(self._md).rstrip() + "\n")
            out.append(md_path)
        return out
